"""Generate a DOCX document with step-by-step testing instructions for the SDLC platform."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

# ── Title ───────────────────────────────────────────────────────────
title = doc.add_heading("SDLC Platform — Testing Steps", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    "This document provides step-by-step instructions to test the deployed SDLC platform, "
    "covering the PM Agent (BRD Generation) module and the Analyst Agent module."
)

# ── Table of Contents (manual) ──────────────────────────────────────
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Prerequisites & Environment Setup",
    "2. Developer Login (SSO Bypass)",
    "3. Project Creation / Selection",
    "4. PM Agent — BRD Generation from Transcript",
    "5. PM Agent — Editing the Generated BRD",
    "6. PM Agent — Push BRD to Confluence & Download",
    "7. Confluence Tab — View Generated BRDs",
    "8. Generate Jira Stories / Epics from BRD",
    "9. Jira Tab — View & Push Stories to Jira",
    "10. Analyst Agent — Knowledge Base Chat",
    "11. Analyst Agent — BRD Generation",
    "12. Automated Sync (Confluence ↔ Jira ↔ Orchestrator)",
    "13. Troubleshooting & Known Issues",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ── Helper ──────────────────────────────────────────────────────────
def add_step(number, text, detail=None):
    p = doc.add_paragraph()
    run = p.add_run(f"Step {number}: ")
    run.bold = True
    p.add_run(text)
    if detail:
        d = doc.add_paragraph(detail)
        d.paragraph_format.left_indent = Inches(0.5)
        d.style.font.size = Pt(10)

def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run("Note: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
    p.add_run(text)

def add_expected(text):
    p = doc.add_paragraph()
    run = p.add_run("Expected Result: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
    p.add_run(text)


# ═══════════════════════════════════════════════════════════════════
# 1. Prerequisites
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("1. Prerequisites & Environment Setup", level=1)

doc.add_paragraph(
    "Before testing, ensure the following are in place:"
)
prereqs = [
    "Deployed URL is accessible (e.g., https://deluxe.siriusai.com)",
    "Backend API is running and reachable from the frontend",
    "AWS Lambda functions are deployed (brd-chat, brd-from-history, brd-generator, requirements-gathering)",
    "Confluence and Jira instances are configured and accessible",
    "Test transcript files ready (.txt, .docx, and/or .pdf format)",
    "Browser: Chrome or Edge (latest version recommended)",
]
for item in prereqs:
    doc.add_paragraph(item, style="List Bullet")

# ═══════════════════════════════════════════════════════════════════
# 2. Developer Login
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("2. Developer Login (SSO Bypass)", level=1)

doc.add_paragraph(
    "During development/testing, a developer bypass mode is available so that "
    "Azure AD SSO is not required. This bypass will be removed in production when "
    "full Azure AD authentication is enabled."
)

add_step(1, "Open the deployed URL in your browser.")
add_step(2, 'On the login page, look for the "Developer Login" option.')
add_step(3, 'Click "Developer Login" or enter the dev bypass token if prompted.')
add_step(4, "You should be redirected to the home/dashboard page.")

add_expected(
    "You are logged in and can see the project dashboard. "
    "The navigation bar shows your user identity or a dev indicator."
)

add_note(
    "If Azure AD SSO is configured and working, you can also test with "
    '"Sign in with Microsoft" which will go through the full SSO flow.'
)

# ═══════════════════════════════════════════════════════════════════
# 3. Project Creation / Selection
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("3. Project Creation / Selection", level=1)

doc.add_paragraph(
    "Each project links to a Jira project and a Confluence space, enabling "
    "bi-directional sync and BRD management."
)

add_step(1, 'Click "Create New Project" on the dashboard (or select an existing project).')
add_step(2, "Fill in the project details:",
         "• Project Name — a descriptive name for the project\n"
         "• Jira Project — select the linked Jira project from the dropdown\n"
         "• Confluence Space — select the linked Confluence space from the dropdown")
add_step(3, 'Click "Create Project".')
add_step(4, "Verify the project appears on the dashboard with the correct Jira and Confluence links.")

add_expected(
    "Project is created and visible on the dashboard. "
    "The Jira project key and Confluence space key are displayed."
)

add_note(
    "If Jira/Confluence dropdowns are empty, verify that the backend has valid "
    "Atlassian credentials configured and that the sync service is running."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 4. PM Agent — BRD Generation from Transcript
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("4. PM Agent — BRD Generation from Transcript", level=1)

doc.add_paragraph(
    "The PM Agent generates a Business Requirements Document (BRD) from uploaded "
    "meeting transcripts. It supports .txt, .docx, and .pdf files. Multiple files "
    "can be uploaded simultaneously and their content will be concatenated."
)

add_step(1, "Open the project you created (or any existing project).")
add_step(2, 'Navigate to the "PM Agent" or "BRD Generation" tab/section.')
add_step(3, 'Click the "Upload Transcript" button or drag-and-drop files into the upload area.')
add_step(4, "Select one or more transcript files:",
         "• Supported formats: .txt, .docx, .pdf\n"
         "• You can select multiple files at once\n"
         "• Files will be listed in the upload area for review before submission")
add_step(5, 'Click "Generate BRD" (or "Submit") to start the generation process.')
add_step(6, "Wait for the BRD generation to complete.",
         "This typically takes 30–90 seconds depending on transcript length. "
         "A loading indicator should be visible during processing.")
add_step(7, "Review the generated BRD displayed on screen.")

add_expected(
    "A complete BRD is generated with sections such as: Executive Summary, "
    "Project Scope, Functional Requirements, Non-Functional Requirements, "
    "Assumptions & Constraints, etc. The content should be derived from the "
    "uploaded transcript(s)."
)

add_note(
    "If multiple files were uploaded, verify that content from ALL files "
    "is reflected in the generated BRD, not just the first file."
)

# ═══════════════════════════════════════════════════════════════════
# 5. PM Agent — Editing the Generated BRD
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("5. PM Agent — Editing the Generated BRD", level=1)

doc.add_paragraph(
    "After BRD generation, users can edit individual sections or the entire BRD."
)

doc.add_heading("5a. Section-wise Editing", level=2)
add_step(1, "In the generated BRD view, click on a specific section heading (e.g., 'Executive Summary').")
add_step(2, "An edit panel or inline editor should appear for that section.")
add_step(3, "Modify the content — add, remove, or rephrase text.")
add_step(4, 'Click "Save" or "Update Section" to save changes.')

add_expected(
    "The section is updated with your changes. Other sections remain unchanged."
)

doc.add_heading("5b. Full BRD Editing", level=2)
add_step(1, 'Click "Edit Entire BRD" or "View Full BRD" to see the complete document.')
add_step(2, "Make edits across any section in the full-document view.")
add_step(3, "Save the changes.")

add_expected(
    "All changes are persisted. Refreshing the page should show the updated BRD."
)

# ═══════════════════════════════════════════════════════════════════
# 6. Push BRD to Confluence & Download
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("6. PM Agent — Push BRD to Confluence & Download", level=1)

add_step(1, 'In the BRD view, click "Push to Confluence".')
add_step(2, "Confirm the target Confluence space (should default to the project's linked space).")
add_step(3, "Wait for the push to complete.")

add_expected(
    "A success message is displayed. The BRD is now available as a Confluence page "
    "in the linked space."
)

add_step(4, '(Optional) Click "Download as DOCX" to download the BRD as a Word document.')

add_expected(
    "A .docx file is downloaded containing the full BRD content with proper formatting."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 7. Confluence Tab
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("7. Confluence Tab — View Generated BRDs", level=1)

add_step(1, 'Navigate to the "Confluence" tab within the project.')
add_step(2, "You should see a list of BRDs that have been pushed to Confluence.")
add_step(3, "Click on a BRD to view its content.")
add_step(4, "Verify the content matches what was generated and edited in the PM Agent.")

add_expected(
    "All pushed BRDs are listed. Clicking on one shows its full content. "
    "Content matches the latest version from the PM Agent."
)

# ═══════════════════════════════════════════════════════════════════
# 8. Generate Jira Stories
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("8. Generate Jira Stories / Epics from BRD", level=1)

add_step(1, 'From the Confluence tab, select a BRD and click "Generate Jira Stories" '
         '(or navigate to the Jira generation section).')
add_step(2, "The system will analyze the BRD and generate Epics and User Stories.")
add_step(3, "Wait for generation to complete (this may take 30–60 seconds).")
add_step(4, "Review the generated Epics and Stories:",
         "• Each Epic should represent a major feature area from the BRD\n"
         "• Stories should be broken down under their respective Epics\n"
         "• Each Story should have a title, description, and acceptance criteria")

add_expected(
    "Epics and Stories are generated that align with the BRD content. "
    "The structure is logical and covers the major requirements from the BRD."
)

# ═══════════════════════════════════════════════════════════════════
# 9. Jira Tab
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("9. Jira Tab — View & Push Stories to Jira", level=1)

add_step(1, 'Navigate to the "Jira" tab within the project.')
add_step(2, "View the list of generated Epics and Stories.")
add_step(3, "Review individual stories — click on a story to see its full details.")
add_step(4, 'Click "Push to Jira" to push the stories to the linked Jira project.')
add_step(5, "Wait for the push to complete.")

add_expected(
    "Stories are pushed to Jira successfully. A confirmation message is shown. "
    "The stories should appear in the linked Jira project with correct "
    "Epic links, descriptions, and acceptance criteria."
)

add_step(6, "Verify in Jira: Open the linked Jira project and confirm the stories exist.")

add_note(
    "If push fails, check that the backend has valid Jira API credentials "
    "and that the target Jira project allows issue creation."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 10. Analyst Agent — Knowledge Base Chat
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("10. Analyst Agent — Knowledge Base Chat", level=1)

doc.add_paragraph(
    "The Analyst Agent uses RAG (Retrieval-Augmented Generation) to answer "
    "questions based on synced Confluence and Jira content for the project."
)

add_step(1, 'Navigate to the "Analyst Agent" tab/section within the project.')
add_step(2, 'Ensure the project has synced data (check the "Sync Docs" status).',
         "If not synced, click 'Sync Docs' and wait for the sync to complete.")
add_step(3, "Type a question in the chat input related to the project's Confluence/Jira content.",
         "Example: 'What are the main requirements for the authentication module?'")
add_step(4, "Submit the query and wait for the response.")

add_expected(
    "The Analyst Agent returns a relevant answer based on the synced knowledge base. "
    "Sources (Confluence pages or Jira issues) are cited in the response."
)

add_step(5, "Verify source citations — click on source links to confirm they point to real Confluence pages or Jira issues.")
add_step(6, "Ask follow-up questions to test the conversational context.")

add_note(
    "The quality of answers depends on the synced content. If answers seem "
    "incomplete, try syncing docs again to pull the latest content."
)

# ═══════════════════════════════════════════════════════════════════
# 11. Analyst Agent — BRD Generation
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("11. Analyst Agent — BRD Generation", level=1)

doc.add_paragraph(
    "The Analyst Agent can also generate BRDs, but instead of using uploaded "
    "transcripts, it leverages the synced Confluence and Jira knowledge base."
)

add_step(1, 'In the Analyst Agent section, look for "Generate BRD" or similar option.')
add_step(2, "Provide a topic or prompt for BRD generation.",
         "Example: 'Generate a BRD for the user authentication and SSO integration feature'")
add_step(3, "Submit and wait for BRD generation (30–90 seconds).")
add_step(4, "Review the generated BRD.")

add_expected(
    "A BRD is generated using context from the project's synced knowledge base. "
    "The content should reference actual project details from Confluence/Jira."
)

add_step(5, "Follow the same editing, push to Confluence, and Jira story generation "
         "steps as described in Sections 5–9 above.")

# ═══════════════════════════════════════════════════════════════════
# 12. Automated Sync
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("12. Automated Sync (Confluence / Jira / Orchestrator)", level=1)

doc.add_paragraph(
    "The platform supports incremental sync between Confluence, Jira, and the "
    "orchestrator's vector store. This keeps the Analyst Agent's knowledge base "
    "up to date."
)

add_step(1, 'Click "Sync Docs" in the project view.')
add_step(2, "The system triggers an incremental sync:",
         "• Fetches changed Confluence pages since last sync\n"
         "• Fetches changed Jira issues since last sync\n"
         "• Updates embeddings in the vector store")
add_step(3, "Wait for the sync to complete. A status indicator should show progress.")
add_step(4, "Verify sync results:",
         "• Check the sync status shows updated counts\n"
         "• Ask the Analyst Agent about recently added/changed content\n"
         "• Confirm new content is reflected in responses")

add_expected(
    "Sync completes successfully. New or modified Confluence pages and Jira issues "
    "are reflected in the Analyst Agent's responses."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 13. Troubleshooting
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("13. Troubleshooting & Known Issues", level=1)

# Table
table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Issue"
hdr_cells[1].text = "Possible Cause"
hdr_cells[2].text = "Resolution"

issues = [
    (
        "401 Unauthorized on API calls",
        "Dev bypass token expired or not set; Azure AD token expired",
        "Re-login using Developer Login or refresh the page. Check browser console for token errors.",
    ),
    (
        "BRD generation fails or times out",
        "Lambda function cold start; large transcript",
        "Wait and retry. Check Lambda logs in CloudWatch for errors. Ensure Lambda runtime is Python 3.12.",
    ),
    (
        "Push to Confluence fails",
        "Invalid Confluence credentials; space permissions",
        "Verify Confluence API token in backend config. Ensure the user has write access to the space.",
    ),
    (
        "Push to Jira fails",
        "Invalid Jira credentials; project permissions",
        "Verify Jira API token in backend config. Ensure the user can create issues in the project.",
    ),
    (
        "Sync Docs shows error",
        "Network issue; Atlassian API rate limit",
        "Retry after a few minutes. Check backend logs for specific error messages.",
    ),
    (
        "Analyst Agent gives irrelevant answers",
        "Knowledge base not synced or incomplete",
        "Run Sync Docs to update the knowledge base. Verify Confluence/Jira content exists.",
    ),
    (
        "File upload rejected",
        "Unsupported file format",
        "Only .txt, .docx, and .pdf files are supported. Ensure the file extension is correct.",
    ),
    (
        "Multiple files not all reflected in BRD",
        "Backend not processing all files",
        "Check that all files were uploaded (listed in upload area). Check backend logs for file processing.",
    ),
]

for issue, cause, resolution in issues:
    row_cells = table.add_row().cells
    row_cells[0].text = issue
    row_cells[1].text = cause
    row_cells[2].text = resolution

# Set column widths
for row in table.rows:
    row.cells[0].width = Inches(2.0)
    row.cells[1].width = Inches(2.5)
    row.cells[2].width = Inches(2.5)

doc.add_paragraph()  # spacing

doc.add_heading("Test Completion Checklist", level=2)

checklist = [
    "Developer login works (SSO bypass or Azure AD)",
    "Project creation with Jira/Confluence linking works",
    "Single file upload (.txt) — BRD generated successfully",
    "Single file upload (.docx) — BRD generated successfully",
    "Single file upload (.pdf) — BRD generated successfully",
    "Multiple file upload — all content reflected in BRD",
    "Section-wise BRD editing works",
    "Full BRD editing works",
    "Push BRD to Confluence — page created in correct space",
    "Download BRD as DOCX — file downloads with correct content",
    "Generate Jira Stories from BRD — Epics and Stories created",
    "Push Stories to Jira — issues appear in Jira project",
    "Analyst Agent chat — relevant answers with source citations",
    "Analyst Agent BRD generation — BRD based on knowledge base",
    "Sync Docs — incremental sync completes successfully",
]

for item in checklist:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run("[ ] ")
    run.font.name = "Consolas"
    p.add_run(item)

# ── Save ────────────────────────────────────────────────────────────
output_dir = os.path.dirname(os.path.abspath(__file__))
output_path = os.path.join(output_dir, "SDLC_Platform_Testing_Steps.docx")
doc.save(output_path)
print(f"Document saved to: {output_path}")