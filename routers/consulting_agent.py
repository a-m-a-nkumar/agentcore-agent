"""
FastAPI router for Joseph — the AI Use Case Prioritization Consultant.

Endpoints
---------
POST /chat       — Server-Sent-Events stream. Mixes text chunks (type=chunk)
                   and structured state updates (type=state) on a single
                   stream so the UI side panel updates live as tools fire.
POST /upload     — Multipart file upload; parses PDF/DOCX/TXT and stashes
                   text against a file_id the agent can later ingest.
GET  /siblings   — Returns the pre-canned sibling assessments for the
                   2x2 matrix portfolio view.
POST /export     — Renders a markdown report as PDF or DOCX.

All endpoints require the existing Azure-AD-backed `get_current_user`
dependency, mirroring `routers/orchestration.py`.
"""

from __future__ import annotations

import asyncio
import io
import json
import logging
import os
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile
from fastapi.responses import StreamingResponse, Response
from pydantic import BaseModel

# Strands + LLM gateway
try:
    from strands import Agent
    from strands.models.openai import OpenAIModel
    from strands.models import BedrockModel
    _STRANDS_AVAILABLE = True
except ImportError as e:
    _STRANDS_AVAILABLE = False
    logging.getLogger(__name__).warning("Strands not available: %s", e)

# Local imports
import re as _re
from consulting_agent_tools import (
    CONSULTING_AGENT_TOOLS,
    fetch_confluence as _fetch_confluence_plain,
    fetch_jira as _fetch_jira_plain,
)
from joseph_prompt import get_full_prompt
from services.consulting_state import (
    KBSnapshot,
    SUB_SCORE_KEYS,
    UploadedFile,
    clear_request_context,
    get_or_create_state,
    get_state,
    push_state_event,
    reset_state,
    set_request_context,
)

# Human-readable labels for the six sub-scores, used when telling Joseph which
# rationale the user edited in the scoring panel.
_SUBSCORE_LABELS = {
    "financial": "Financial Impact",
    "productivity": "Scale of Impact on Productivity",
    "intent": "Business Intent and Need",
    "complexity": "Implementation Complexity",
    "data_platform": "Data and Platform Readiness",
    "measurement": "Ease of Measuring Success",
}
from services.knowledge_base import (
    fetch_full_content as kb_fetch_full_content,
    lookup_by_id_or_title as kb_lookup,
    search as kb_search,
)
from services.marker_parser import parse_and_fire_events

# Regexes for server-side preprocessing of user messages
_CONFLUENCE_URL_RE = _re.compile(
    r"https?://[\w.-]+\.atlassian\.net/(?:wiki/)?(?:[\w-]+/)*pages/(\d+)[\w/?&=#-]*",
    _re.IGNORECASE,
)
_JIRA_KEY_RE = _re.compile(
    r"(?:https?://[\w.-]+\.atlassian\.net/browse/)?([A-Z][A-Z0-9]+-\d+)\b"
)
# "Consume document 01", "Read kb-002", "Consume the ResponseAI proposal"
_CONSUME_INTENT_RE = _re.compile(
    r"(?:consume|read|open|pull up|load|fetch|ingest)\s+(?:document\s+|doc\s+|the\s+)?[\"']?(kb-\d+|\d{1,3}|[A-Za-z][\w\s\-—:]{3,80}?)[\"']?(?=[,.;\n]|$|\s+and\s+|\s+plus\s+)",
    _re.IGNORECASE,
)

# Auth — reuse the existing dependency from app.py (it's defined in app.py
# but conventionally available via `routers.orchestration.get_current_user`).
from routers.orchestration import get_current_user

logger = logging.getLogger(__name__)
router = APIRouter()

MOCK_DIR = Path(__file__).parent.parent / "mock_data"

# Optional AgentCore Memory — gracefully degrade to a process-local fallback
# so the POC runs on a laptop without AWS creds. We mirror app.py's boto3
# client setup (timeouts + retries) so credential resolution is identical to
# the existing analyst-history endpoint.
try:
    from environment import DEFAULT_AGENTCORE_MEMORY_ID, DEFAULT_AGENTCORE_ACTOR_ID
    import boto3
    _AGENTCORE_AVAILABLE = bool(DEFAULT_AGENTCORE_MEMORY_ID)
    _AGENTCORE_MEMORY_ID = DEFAULT_AGENTCORE_MEMORY_ID
    _AGENTCORE_ACTOR_ID = DEFAULT_AGENTCORE_ACTOR_ID
except Exception:
    _AGENTCORE_AVAILABLE = False
    _AGENTCORE_MEMORY_ID = None
    _AGENTCORE_ACTOR_ID = None

# Set to False the first time an AgentCore call fails so we stop trying for
# the lifetime of the process and stop spamming the log. This is the same
# "noisy-once, quiet-after" pattern used elsewhere when an external service
# is intermittently unreachable.
_AGENTCORE_RUNTIME_DISABLED = False

_FALLBACK_HISTORY: dict[str, list[dict[str, str]]] = {}

# In-process Insights store. Each entry is a completed assessment that the
# user explicitly saved from the agent. Restart loses these — no DB by
# design for the POC. Pre-loaded siblings are merged with these at read time.
_INSIGHTS_SAVED: list[dict[str, Any]] = []


# ───────────────────────── Request / response models ─────────────────────────

class ChatRequest(BaseModel):
    message: str
    session_id: Optional[str] = None
    reset: bool = False


class RescoreRequest(BaseModel):
    session_id: str
    sub_score: str
    rationale: str


class ExportRequest(BaseModel):
    markdown_report: str
    format: str = "pdf"
    filename: Optional[str] = None


class SaveToInsightsRequest(BaseModel):
    session_id: str
    title: str
    sponsor: Optional[str] = None
    report_markdown: str


# ───────────────────────── Memory helpers ─────────────────────────

def _get_memory_client():
    """Build a bedrock-agentcore client using the same timeout/retry config
    as `app.get_agent_core_client()` so credential resolution is identical."""
    from botocore.config import Config
    region = os.getenv("BEDROCK_REGION") or os.getenv("AWS_REGION") or "us-east-1"
    config = Config(read_timeout=300, connect_timeout=10, retries={"max_attempts": 3})
    return boto3.client("bedrock-agentcore", region_name=region, config=config)


def _agentcore_usable() -> bool:
    return _AGENTCORE_AVAILABLE and not _AGENTCORE_RUNTIME_DISABLED


def _disable_agentcore_runtime(reason: str) -> None:
    """Called once on the first AgentCore failure. Flips a module-level
    flag so subsequent calls skip the boto3 attempt and go straight to the
    in-process fallback — avoids spamming the log every turn."""
    global _AGENTCORE_RUNTIME_DISABLED
    if _AGENTCORE_RUNTIME_DISABLED:
        return
    _AGENTCORE_RUNTIME_DISABLED = True
    logger.info(
        "AgentCore Memory disabled for this process: %s. "
        "Falling back to in-process conversation history. "
        "Restart with valid AWS creds (AWS_PROFILE or AWS_ACCESS_KEY_ID) to re-enable.",
        reason,
    )


def _append_to_memory(session_id: str, role: str, content: str) -> None:
    if _agentcore_usable():
        try:
            client = _get_memory_client()
            role_map = {"user": "USER", "assistant": "ASSISTANT", "system": "OTHER"}
            client.create_event(
                memoryId=_AGENTCORE_MEMORY_ID,
                actorId=_AGENTCORE_ACTOR_ID,
                sessionId=session_id,
                eventTimestamp=datetime.utcnow(),
                payload=[{
                    "conversational": {
                        "role": role_map.get(role, "OTHER"),
                        "content": {"text": (content or "").strip()[:9000] or "(empty)"},
                    }
                }],
                clientToken=str(uuid.uuid4()),
            )
            return
        except Exception as e:
            _disable_agentcore_runtime(str(e).split(":")[0] or "unknown error")
    _FALLBACK_HISTORY.setdefault(session_id, []).append({"role": role, "content": content})


def _load_history(session_id: str, max_messages: int = 30) -> list[dict[str, str]]:
    if _agentcore_usable():
        try:
            client = _get_memory_client()
            response = client.list_events(
                memoryId=_AGENTCORE_MEMORY_ID,
                sessionId=session_id,
                actorId=_AGENTCORE_ACTOR_ID,
                includePayloads=True,
                maxResults=max_messages,
            )
            events = response.get("events", [])
            events_sorted = sorted(
                events,
                key=lambda e: (
                    e.get("eventTimestamp").isoformat()
                    if hasattr(e.get("eventTimestamp"), "isoformat")
                    else str(e.get("eventTimestamp", ""))
                ),
            )
            messages = []
            for ev in events_sorted:
                for p in ev.get("payload", []):
                    conv = p.get("conversational") or {}
                    text = conv.get("content", {}).get("text")
                    role_raw = conv.get("role", "OTHER")
                    role = {"USER": "user", "ASSISTANT": "assistant"}.get(role_raw, "system")
                    if text:
                        messages.append({"role": role, "content": text})
            return messages
        except Exception as e:
            _disable_agentcore_runtime(str(e).split(":")[0] or "unknown error")
    return list(_FALLBACK_HISTORY.get(session_id, []))


# ───────────────────────── Agent construction ─────────────────────────

def _build_agent():
    """Build a fresh Strands agent for one /chat request.

    Joseph forces `BedrockModel` (direct Bedrock) instead of `OpenAIModel`
    (DLX gateway). The DLX gateway proxies OpenAI tool-use protocol to
    Claude's native tool-use protocol, and that translation 500s on the
    `tool_result` follow-up call. Going direct to Bedrock uses Claude's
    native protocol end-to-end, so tool round-trips work — same model
    (claude-sonnet-4-5), same account, just no broken translator.

    Set CONSULTING_USE_GATEWAY=1 in the env to override this and use the
    gateway anyway (useful when running outside AWS, but expect tool
    follow-ups to fail).
    """
    if not _STRANDS_AVAILABLE:
        raise RuntimeError("Strands SDK is not installed.")

    try:
        from environment import (
            AGENT_MODEL_PROVIDER,
            DEFAULT_DLXAI_GATEWAY_URL,
            DEFAULT_DLXAI_GATEWAY_KEY,
            DEFAULT_GATEWAY_MODEL,
        )
    except ImportError:
        AGENT_MODEL_PROVIDER = "openai"
        DEFAULT_DLXAI_GATEWAY_URL = "https://dlxai-dev.deluxe.com/proxy"
        DEFAULT_DLXAI_GATEWAY_KEY = ""
        DEFAULT_GATEWAY_MODEL = "Claude-4.5-Sonnet"

    # Use the DLX gateway with NO tools attached. Joseph emits structured
    # `[[JOSEPH_EVENT:...]]` markers inline in his prose; the runtime
    # parses them after the agent returns and pushes SSE state events.
    # This avoids the gateway's broken tool-result translation entirely
    # (one model call per turn, no tool follow-ups).
    #
    # We previously tried direct Bedrock (BedrockModel) but the
    # organization's SCP explicitly denies bedrock:InvokeModel* on user
    # SSO roles. Lambdas have a different IAM context and can use Bedrock
    # directly — see analyst_agent's Lambda for that pattern.
    gateway_url = os.getenv("DLXAI_GATEWAY_URL", DEFAULT_DLXAI_GATEWAY_URL)
    gateway_key = os.getenv("DLXAI_GATEWAY_KEY", DEFAULT_DLXAI_GATEWAY_KEY)
    gateway_model = os.getenv("GATEWAY_MODEL", DEFAULT_GATEWAY_MODEL)
    logger.info("Joseph: using OpenAIModel via DLX gateway, tools=[] (model=%s)", gateway_model)
    model = OpenAIModel(
        model_id=gateway_model,
        client_args={"base_url": gateway_url, "api_key": gateway_key},
    )

    return Agent(
        model=model,
        tools=[],
        system_prompt=get_full_prompt(),
    )


def _build_input(history: list[dict[str, str]], user_message: str, session_id: str) -> str:
    """Compose the agent input. Prior turns are prepended as context so a
    fresh Agent instance still sees prior conversation. Session id is
    surfaced so Joseph can pass it through to tools per the runtime addendum.
    """
    parts = [f"[Session ID for tool calls]: {session_id}", ""]
    if history:
        parts.append("[Prior conversation]")
        for m in history[-20:]:
            role = "User" if m["role"] == "user" else "Joseph"
            parts.append(f"{role}: {m['content']}")
        parts.append("")
    parts.append("[Current user message]")
    parts.append(user_message)
    return "\n".join(parts)


def _extract_text(result: Any) -> str:
    """
    Pull the plain text out of a Strands AgentResult.

    AgentResult.message is typically an OpenAI-style dict:
        {"role": "assistant", "content": [{"text": "...prose..."}]}
    Falling back to str(result) on that gives a Python dict repr with
    ESCAPED newlines ("\\n") which then breaks JSON parsing of any
    [[JOSEPH_EVENT:...]] markers inside the text. So we walk the
    structure explicitly and concatenate the actual text fields.
    """
    if result is None:
        return ""
    if isinstance(result, str):
        return result

    msg = getattr(result, "message", None)

    # Most common shape: AgentResult.message is a dict with content list
    if isinstance(msg, dict):
        content = msg.get("content")
        if isinstance(content, list):
            parts: list[str] = []
            for item in content:
                if isinstance(item, dict):
                    t = item.get("text")
                    if t:
                        parts.append(str(t))
                elif isinstance(item, str):
                    parts.append(item)
            if parts:
                return "\n".join(parts)
        if isinstance(content, str):
            return content
    if isinstance(msg, str):
        return msg

    for attr in ("data", "output", "text", "content"):
        v = getattr(result, attr, None)
        if v is None:
            continue
        if isinstance(v, str):
            return v
        if isinstance(v, list):
            parts = []
            for item in v:
                if isinstance(item, dict) and "text" in item:
                    parts.append(str(item["text"]))
                elif isinstance(item, str):
                    parts.append(item)
            if parts:
                return "\n".join(parts)
        if isinstance(v, dict) and "text" in v:
            return str(v["text"])

    return str(result)


def _preprocess_user_message(session_id: str, message: str) -> str:
    """
    Server-side enrichment that runs before the agent sees the user message.

    Order of operations:
      1. Auto-search the internal KB on the FIRST inquiry of the session
         (when `kb_search_done` is False). Push the results to the UI via
         a state event and inject a summary of hits into the agent prompt
         as context — but do NOT inline full content yet. The user picks
         which to consume.
      2. Detect "consume / read / open <id-or-title>" intent in this
         message → look up the matching KB entry → inline its full
         content + mark as consumed.
      3. Inline any unconsumed uploaded files (drag-and-drop path).
      4. Inline any Confluence URLs pasted directly.
      5. Inline any Jira issue keys mentioned.

    All inlining is best-effort. If something fails the message goes
    through unchanged.
    """
    parts: list[str] = []
    state = get_or_create_state(session_id)

    # 1. Auto-search the KB on the first inquiry
    if not state.kb_search_done and len(message.strip()) >= 20:
        try:
            results = kb_search(message, limit=10)
            state.kb_results = [
                KBSnapshot(
                    id=r.id, title=r.title, url=r.url, snippet=r.snippet,
                    type=r.type, icon=r.icon, relevance=r.relevance,
                )
                for r in results
            ]
            state.kb_search_done = True
            push_state_event("kb_results", state.to_kb_payload())

            if results:
                top_lines = []
                for i, r in enumerate(results[:6], start=1):
                    top_lines.append(
                        f"{i:02d}. {r.title}\n"
                        f"    URL: {r.url}\n"
                        f"    Relevance: {r.relevance:.2f}\n"
                        f"    Snippet: {r.snippet[:280]}"
                    )
                parts.append(
                    "[INTERNAL KNOWLEDGE BASE — auto-search results]\n"
                    "The runtime ran a keyword search on the user's message and surfaced these documents.\n"
                    "The UI is showing all 10 hits as cards. Summarize the top 3-5 in your response and\n"
                    "ask the user which they want you to read. Do NOT fabricate detail about a doc you\n"
                    "have not been given the full text of — only what's in the snippet here.\n\n"
                    + "\n\n".join(top_lines)
                    + "\n[END KNOWLEDGE BASE RESULTS]"
                )
        except Exception as e:
            logger.warning("KB auto-search failed: %s", e)

    # 2. Detect consume intent and inline full KB content
    consumed_this_turn: set[str] = set()
    for m in _CONSUME_INTENT_RE.finditer(message):
        needle = (m.group(1) or "").strip()
        if not needle:
            continue
        entry = kb_lookup(needle)
        if not entry:
            continue
        kb_id = entry["id"]
        if kb_id in state.consumed_kb_ids or kb_id in consumed_this_turn:
            continue
        fetched = kb_fetch_full_content(kb_id)
        if not fetched:
            continue
        _, full_text = fetched
        snippet = full_text if len(full_text) <= 12000 else full_text[:12000] + "\n\n[...truncated]"
        parts.append(
            f"[KB DOCUMENT {kb_id} · {entry['title']}]\n"
            f"URL: {entry['url']}\n\n"
            f"{snippet}\n"
            f"[END KB DOCUMENT]"
        )
        state.consumed_kb_ids.add(kb_id)
        consumed_this_turn.add(kb_id)

    if consumed_this_turn:
        push_state_event("kb_results", state.to_kb_payload())

    # 3. Uploaded files (drag-and-drop path — unchanged)
    unread = [f for f in state.uploaded_files.values() if not f.consumed]
    for f in unread:
        snippet = f.text if len(f.text) <= 12000 else f.text[:12000] + "\n\n[...truncated]"
        parts.append(f"[DOCUMENT: {f.filename}]\n{snippet}\n[END DOCUMENT]")
        f.consumed = True

    # 4. Confluence URLs (unchanged)
    for m in _CONFLUENCE_URL_RE.finditer(message):
        page_id = m.group(1)
        try:
            content = _fetch_confluence_plain(page_id)
            parts.append(f"[CONFLUENCE PAGE {page_id}]\n{content}\n[END CONFLUENCE PAGE]")
        except Exception as e:
            logger.warning("Confluence inline-fetch failed for page %s: %s", page_id, e)

    # 5. Jira issue keys (unchanged)
    seen_keys: set[str] = set()
    for m in _JIRA_KEY_RE.finditer(message):
        key = m.group(1).upper()
        if key in seen_keys:
            continue
        seen_keys.add(key)
        try:
            content = _fetch_jira_plain(key)
            parts.append(f"[JIRA ISSUE {key}]\n{content}\n[END JIRA ISSUE]")
        except Exception as e:
            logger.warning("Jira inline-fetch failed for %s: %s", key, e)

    if not parts:
        return message
    return "\n\n".join(parts) + "\n\n---\n\n[USER MESSAGE]\n" + message


# ───────────────────────── /chat ─────────────────────────

@router.post("/chat")
async def chat(
    request: ChatRequest,
    current_user: dict = Depends(get_current_user),
):
    user_id = current_user.get("user_id") or current_user.get("oid") or "anonymous"

    if request.reset or not request.session_id:
        session_id = f"joseph-session-{user_id}-{uuid.uuid4()}"
        reset_state(session_id)
        _FALLBACK_HISTORY.pop(session_id, None)
    else:
        session_id = request.session_id
        get_or_create_state(session_id)

    queue: asyncio.Queue = asyncio.Queue()

    async def generate_sse():
        # Emit metadata first so frontend can store the session id.
        yield f"data: {json.dumps({'type': 'metadata', 'session_id': session_id})}\n\n"

        if not _STRANDS_AVAILABLE:
            yield f"data: {json.dumps({'type': 'error', 'message': 'Strands SDK not installed'})}\n\n"
            yield f"data: {json.dumps({'type': 'done'})}\n\n"
            return

        set_request_context(session_id, queue)
        history = _load_history(session_id)
        _append_to_memory(session_id, "user", request.message)
        enriched_user_message = _preprocess_user_message(session_id, request.message)
        prompt_input = _build_input(history, enriched_user_message, session_id)

        async def _run_agent():
            try:
                agent = _build_agent()
                loop = asyncio.get_running_loop()
                result = await loop.run_in_executor(None, lambda: agent(prompt_input))
                final_text = _extract_text(result).strip()
                await queue.put({"type": "_final", "text": final_text})
            except Exception as e:
                logger.exception("Agent run failed")
                await queue.put({"type": "error", "message": f"Agent error: {e}"})
            finally:
                await queue.put({"type": "_complete"})

        agent_task = asyncio.create_task(_run_agent())

        final_text: str | None = None
        try:
            while True:
                event = await queue.get()
                etype = event.get("type")
                if etype == "_complete":
                    break
                if etype == "_final":
                    final_text = event.get("text", "")
                    continue
                if etype == "state":
                    yield f"data: {json.dumps(event)}\n\n"
                elif etype == "error":
                    yield f"data: {json.dumps(event)}\n\n"

            if final_text:
                # Parse any [[JOSEPH_EVENT:...]] markers, mutate state,
                # push SSE state events for scoring/coverage/citations,
                # then return the text with markers stripped.
                cleaned_text = parse_and_fire_events(session_id, final_text)

                # Re-emit the latest state snapshots so the UI gets the
                # update right before the prose starts streaming. (The
                # marker parser already pushed events to the queue, but
                # we already drained the queue above — emit explicitly.)
                state = get_or_create_state(session_id)
                yield (
                    f"data: {json.dumps({'type': 'state', 'kind': 'scores', 'payload': state.to_scores_payload()})}\n\n"
                )
                yield (
                    f"data: {json.dumps({'type': 'state', 'kind': 'coverage', 'payload': state.to_coverage_payload()})}\n\n"
                )
                if state.citations:
                    yield (
                        f"data: {json.dumps({'type': 'state', 'kind': 'citations', 'payload': state.to_citations_payload()})}\n\n"
                    )

                _append_to_memory(session_id, "assistant", cleaned_text)

                # Final feasibility report heuristic — fire a report event
                # so the UI surfaces the export buttons.
                if "# Use Case Feasibility Report" in cleaned_text:
                    state.current_report = cleaned_text
                    yield (
                        f"data: {json.dumps({'type': 'state', 'kind': 'report', 'payload': cleaned_text})}\n\n"
                    )

                # Pseudo-stream the cleaned text in 3-word chunks.
                words = cleaned_text.split(" ")
                chunk = ""
                for i, word in enumerate(words):
                    chunk += word + " "
                    if (i + 1) % 3 == 0 or i == len(words) - 1:
                        yield f"data: {json.dumps({'type': 'chunk', 'content': chunk})}\n\n"
                        await asyncio.sleep(0.02)
                        chunk = ""

            yield f"data: {json.dumps({'type': 'done'})}\n\n"
        finally:
            clear_request_context()
            if not agent_task.done():
                agent_task.cancel()

    return StreamingResponse(
        generate_sse(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


# ───────────────────────── /rescore ─────────────────────────

@router.post("/rescore")
async def rescore(
    request: RescoreRequest,
    current_user: dict = Depends(get_current_user),
):
    """Re-evaluate a single sub-score after the user edits its rationale in the
    scoring panel.

    The user's edited rationale is treated as new evidence. Joseph re-scores
    that sub-score (and any other sub-score the new information genuinely
    bears on), re-emits the full `[[JOSEPH_EVENT:scores]]` block, and we return
    the updated scores payload. The axis averages and quadrant are computed
    properties of the state, so the overall placement updates automatically.

    Non-streaming: a rationale edit is a small, focused round-trip, so we run
    the agent once and return JSON rather than opening an SSE stream.
    """
    if not _STRANDS_AVAILABLE:
        raise HTTPException(status_code=503, detail="Strands SDK not installed")

    key = request.sub_score
    if key not in SUB_SCORE_KEYS:
        raise HTTPException(
            status_code=400,
            detail=f"Unknown sub_score '{key}'. Valid: {list(SUB_SCORE_KEYS)}",
        )

    new_rationale = (request.rationale or "").strip()
    if not new_rationale:
        raise HTTPException(status_code=400, detail="rationale cannot be empty")

    state = get_state(request.session_id)
    if state is None:
        raise HTTPException(
            status_code=404,
            detail=f"No active session state for {request.session_id}",
        )

    session_id = request.session_id
    label = _SUBSCORE_LABELS.get(key, key)

    # Persist the user's edit immediately as the "consumed" facts of record, so
    # it isn't lost even if the agent call fails. The user edits the inputs;
    # Joseph re-derives the `ranking` (why-this-level) judgement on re-score.
    state.scores[key].consumed = new_rationale

    # Snapshot current scores so Joseph knows his own baseline — markers are
    # stripped from stored history, so prior numeric scores aren't in context.
    snapshot_lines = []
    for k in SUB_SCORE_KEYS:
        ss = state.scores[k]
        v = ss.value if ss.value is not None else "—"
        c = ss.confidence or "—"
        snapshot_lines.append(f"- {_SUBSCORE_LABELS[k]} ({k}): {v}/5, confidence {c}")
    current_snapshot = "\n".join(snapshot_lines)

    instruction = (
        "[FACTS EDIT — re-score request]\n"
        f'The user opened the scoring panel and edited the "Consumed" facts for the '
        f'"{label}" sub-score ({key}) — the inputs the score is built from. Their '
        "revised / added facts are:\n\n"
        f'"""\n{new_rationale}\n"""\n\n'
        "Treat this as corrected evidence from the user. Re-evaluate the 1-5 value and "
        f"confidence for {key} against the scoring framework, using these facts. "
        "If — and only if — they also bear on another sub-score, adjust that one too; "
        "otherwise leave the others exactly as they are.\n\n"
        "Current scores on record:\n"
        f"{current_snapshot}\n\n"
        "Re-emit the FULL [[JOSEPH_EVENT:scores]] block with all six sub-scores "
        "(keep unchanged ones at their current value/confidence and both text "
        "fields). For the edited sub-score: `consumed` must preserve the substance "
        "of the facts the user wrote (tightened into clean prose, not discarded), and "
        "`ranking` is your own 2-3 sentence justification — re-derived from those "
        "facts — for why it lands at this band.\n\n"
        "Then reply to the user in 1-2 sentences: say whether the score moved, to what, "
        "and why. Do not run discovery questions or restate the whole framework."
    )

    history = _load_history(session_id)
    prompt_input = _build_input(history, instruction, session_id)

    # No SSE queue here: parse_and_fire_events mutates state directly, and
    # push_state_event no-ops when no queue is set. We read the state back
    # and return it as JSON.
    try:
        agent = _build_agent()
        loop = asyncio.get_running_loop()
        result = await loop.run_in_executor(None, lambda: agent(prompt_input))
        final_text = _extract_text(result).strip()
    except Exception as e:
        logger.exception("Rescore agent run failed")
        # The user's edited rationale is already persisted; surface a clear
        # error but keep the saved text so the panel reflects the edit.
        raise HTTPException(status_code=502, detail=f"Agent error during rescore: {e}")

    cleaned_text = parse_and_fire_events(session_id, final_text)

    # Keep the conversation log coherent for subsequent turns.
    _append_to_memory(
        session_id, "user",
        f'[Edited rationale for "{label}" in the scoring panel]\n{new_rationale}',
    )
    if cleaned_text:
        _append_to_memory(session_id, "assistant", cleaned_text)

    state = get_or_create_state(session_id)
    return {
        "scores": state.to_scores_payload(),
        "message": cleaned_text,
    }


# ───────────────────────── /upload ─────────────────────────

@router.post("/upload")
async def upload(
    session_id: str,
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user),
):
    filename = (file.filename or "upload").strip()
    lower = filename.lower()
    raw = await file.read()
    text = ""

    try:
        if lower.endswith(".pdf"):
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(raw))
            text = "\n\n".join((p.extract_text() or "") for p in reader.pages)
        elif lower.endswith(".docx"):
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(raw))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        elif lower.endswith((".txt", ".md")):
            text = raw.decode("utf-8", errors="replace")
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {filename}")
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Upload parse failed")
        raise HTTPException(status_code=400, detail=f"Could not parse file: {e}")

    file_id = str(uuid.uuid4())
    state = get_or_create_state(session_id)
    state.uploaded_files[file_id] = UploadedFile(
        file_id=file_id, filename=filename, text=text, chars=len(text)
    )

    return {"file_id": file_id, "filename": filename, "chars": len(text)}


# ───────────────────────── /siblings ─────────────────────────

@router.get("/siblings")
async def siblings(current_user: dict = Depends(get_current_user)):
    """Pre-loaded portfolio assessments. Used by the Consulting Agent page
    to populate the matrix sibling dots. Does NOT include in-process saved
    assessments — that's a separate endpoint for the Insights page."""
    path = MOCK_DIR / "sibling_assessments.json"
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


# ───────────────────────── /insights ─────────────────────────

@router.get("/insights")
async def insights(current_user: dict = Depends(get_current_user)):
    """Merged portfolio for the Insights page: pre-loaded siblings + any
    assessments the user saved from the agent during this server process."""
    path = MOCK_DIR / "sibling_assessments.json"
    with open(path, "r", encoding="utf-8") as f:
        siblings_data = json.load(f)
    saved = sorted(
        _INSIGHTS_SAVED,
        key=lambda a: a.get("assessed_at", ""),
        reverse=True,
    )
    # Newest saved first, then siblings — index numbers in the UI follow this order.
    return saved + siblings_data


@router.post("/insights/save")
async def save_to_insights(
    request: SaveToInsightsRequest,
    current_user: dict = Depends(get_current_user),
):
    """Persist the current session's assessment into the in-process Insights
    store. Idempotent on session_id — re-saving updates the existing entry
    rather than adding a duplicate. Lost on server restart."""
    state = get_state(request.session_id)
    if state is None:
        raise HTTPException(
            status_code=404,
            detail=f"No active session state for {request.session_id}",
        )

    axes = state.axes
    if axes["impact"] is None or axes["speed"] is None:
        raise HTTPException(
            status_code=400,
            detail="Assessment isn't scored yet — finish the conversation first",
        )

    quadrant = state.quadrant or "Pending"
    sub_scores = {
        k: int(ss.value) if ss.value is not None else 0
        for k, ss in state.scores.items()
    }

    record = {
        "id": f"session-{request.session_id}",
        "title": request.title,
        "sponsor": request.sponsor or "Self-reported",
        "quadrant": quadrant,
        "axes": {
            "impact": round(axes["impact"], 2),
            "speed": round(axes["speed"], 2),
        },
        "sub_scores": sub_scores,
        "status": "Just assessed",
        "one_liner": (request.title or "")[:200],
        "assessed_at": datetime.utcnow().isoformat(timespec="seconds") + "Z",
        "assessed_by": "Joseph",
        "report_markdown": request.report_markdown,
        "is_fresh": True,
    }

    # Idempotency: replace any existing record with the same session id.
    _INSIGHTS_SAVED[:] = [
        a for a in _INSIGHTS_SAVED if a.get("id") != record["id"]
    ] + [record]

    return {"saved": True, "id": record["id"]}


# ───────────────────────── /export ─────────────────────────

@router.post("/export")
async def export(
    request: ExportRequest,
    current_user: dict = Depends(get_current_user),
):
    fmt = (request.format or "pdf").lower()
    base_name = (request.filename or "feasibility-report").rsplit(".", 1)[0]

    if fmt == "docx":
        try:
            from docx import Document as DocxDocument
            from docx.shared import Pt
        except ImportError:
            raise HTTPException(status_code=500, detail="python-docx not installed")

        doc = DocxDocument()
        for line in request.markdown_report.splitlines():
            stripped = line.strip()
            if stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("- "):
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(stripped[2:])
            elif stripped:
                p = doc.add_paragraph(stripped)
                p.style.font.size = Pt(11)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return Response(
            content=buf.read(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{base_name}.docx"'},
        )

    # PDF — use reportlab if available; otherwise return text/plain with
    # a clear message rather than crashing the demo.
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except ImportError:
        return Response(
            content=request.markdown_report.encode("utf-8"),
            media_type="text/plain",
            headers={
                "Content-Disposition": f'attachment; filename="{base_name}.md"',
                "X-Export-Notice": "reportlab not installed — returning markdown",
            },
        )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=LETTER, topMargin=48, bottomMargin=48)
    styles = getSampleStyleSheet()
    story = []
    for line in request.markdown_report.splitlines():
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 8))
            continue
        if stripped.startswith("# "):
            story.append(Paragraph(stripped[2:], styles["Title"]))
        elif stripped.startswith("## "):
            story.append(Paragraph(stripped[3:], styles["Heading2"]))
        elif stripped.startswith("### "):
            story.append(Paragraph(stripped[4:], styles["Heading3"]))
        else:
            safe = (stripped.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))
            story.append(Paragraph(safe, styles["BodyText"]))
    doc.build(story)
    buf.seek(0)
    return Response(
        content=buf.read(),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{base_name}.pdf"'},
    )
