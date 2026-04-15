"""
Terraform Generator Router
Consumes a SAD (Solution Architecture Document), extracts infrastructure components,
and generates modular Terraform code using Claude (Bedrock).
"""

import base64
import io
import json
import logging
import re
import zipfile
from typing import Optional

import boto3
import httpx
from botocore.config import Config as BotoConfig
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from auth import verify_azure_token

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/terraform", tags=["terraform"])

# ─── Bedrock ──────────────────────────────────────────────────────────────────

_bedrock = None

def get_bedrock():
    global _bedrock
    if _bedrock is None:
        _bedrock = boto3.client(
            "bedrock-runtime",
            region_name="us-east-1",
            config=BotoConfig(read_timeout=300, connect_timeout=10, retries={"max_attempts": 2}),
        )
    return _bedrock

MODEL_ID = "us.anthropic.claude-sonnet-4-5-20250929-v1:0"

def invoke_claude(prompt: str, max_tokens: int = 8000, temperature: float = 1.0) -> str:
    body = json.dumps({
        "anthropic_version": "bedrock-2023-05-31",
        "max_tokens": max_tokens,
        "temperature": temperature,
        "messages": [{"role": "user", "content": prompt}],
    })
    resp = get_bedrock().invoke_model(modelId=MODEL_ID, body=body)
    result = json.loads(resp["body"].read())
    return result["content"][0]["text"]

# ─── Models ───────────────────────────────────────────────────────────────────

class ExtractRequest(BaseModel):
    document_content: str
    document_title: Optional[str] = "Architecture Document"

class ValidateComponentRequest(BaseModel):
    name: str

class Component(BaseModel):
    id: str
    name: str
    description: str
    category: str       # networking, compute, storage, database, security, integration

class GenerateRequest(BaseModel):
    document_content: str
    selected_components: list[Component]
    project_name: Optional[str] = "my-infrastructure"
    aws_region: Optional[str] = "us-east-1"
    environment: Optional[str] = "dev"

class DownloadRequest(BaseModel):
    files: dict[str, str]   # filename -> content
    project_name: Optional[str] = "terraform"

class GitHubPushRequest(BaseModel):
    files: dict[str, str]           # filename -> content
    github_token: str               # GitHub PAT (ghp_...)
    repo_name: str                  # e.g. "my-infra-terraform"
    branch: Optional[str] = "main"
    commit_message: Optional[str] = "feat: add Terraform infrastructure code"
    private: Optional[bool] = True
    folder_prefix: Optional[str] = ""  # optional subfolder inside repo

class HarnessPushRequest(BaseModel):
    files: dict[str, str]           # filename -> content
    api_key: str                    # Harness API key
    repo_url: str                   # Full Harness Code repo URL (browser address bar)
    branch: Optional[str] = "main"
    commit_message: Optional[str] = "feat: add Terraform infrastructure code"
    folder_prefix: Optional[str] = ""

# ─── Helpers ──────────────────────────────────────────────────────────────────

def strip_html(text: str) -> str:
    return re.sub(r'<[^>]+>', ' ', text).strip()

def strip_fences(text: str, lang: str = "") -> str:
    """Remove markdown code fences from generated content."""
    text = text.strip()
    pattern = rf"^```{lang}\s*\n?(.*?)\n?```$"
    m = re.match(pattern, text, re.DOTALL | re.IGNORECASE)
    if m:
        return m.group(1).strip()
    if text.startswith("```"):
        lines = text.split("\n")
        lines = lines[1:]
        if lines and lines[-1].strip() == "```":
            lines = lines[:-1]
        return "\n".join(lines).strip()
    return text


def parse_delimited_files(raw: str) -> dict:
    """
    Parse delimiter-separated Terraform files from Claude's response.
    Format:
        ===FILE: path/to/file.tf===
        <content>
        ===FILE: path/to/other.tf===
        <content>
        ===END===
    Falls back to code-fence parsing if delimiters not found.
    """
    result = {}
    # Split on ===FILE: ... === markers
    parts = re.split(r'===FILE:\s*(.+?)\s*===', raw)
    # parts[0] = preamble, then alternating: filename, content, filename, content ...
    if len(parts) >= 3:
        i = 1
        while i + 1 < len(parts):
            filename = parts[i].strip()
            content = parts[i + 1]
            # Strip trailing ===END=== or similar
            content = re.sub(r'\s*===END===.*$', '', content, flags=re.DOTALL)
            # Strip any markdown fences the model may have added around the content
            content = strip_fences(content)
            result[filename] = content
            i += 2
        return result

    # Fallback: try to find HCL blocks separated by code fences with filename comments
    # e.g.  ```hcl\n# modules/vpc/main.tf\n...```
    fence_blocks = re.findall(r'```(?:hcl|terraform)?\s*\n(.*?)```', raw, re.DOTALL | re.IGNORECASE)
    for block in fence_blocks:
        lines = block.strip().splitlines()
        if lines and lines[0].lstrip('#').strip().endswith('.tf'):
            path = lines[0].lstrip('#').strip()
            result[path] = "\n".join(lines[1:]).strip()

    if result:
        return result

    raise ValueError(f"Could not parse module files from response. Raw response starts with: {raw[:200]}")

# ─── Endpoints ────────────────────────────────────────────────────────────────

@router.post("/parse-document")
async def parse_document(file: UploadFile = File(...), _=Depends(verify_azure_token)):
    """
    Accept a PDF, DOCX, or TXT file and return extracted plain text.
    """
    filename = (file.filename or "").lower()
    raw_bytes = await file.read()

    try:
        if filename.endswith(".pdf"):
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(raw_bytes))
            pages = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
            text = "\n\n".join(pages)

        elif filename.endswith(".docx"):
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(raw_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            text = "\n".join(paragraphs)

        elif filename.endswith((".txt", ".md")):
            text = raw_bytes.decode("utf-8", errors="replace")

        else:
            raise HTTPException(
                status_code=400,
                detail="Unsupported file type. Please upload a PDF, DOCX, or TXT file."
            )

        if not text.strip():
            raise HTTPException(status_code=400, detail="Could not extract any text from the file.")

        return {"text": text, "filename": file.filename, "chars": len(text)}

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to parse file: {str(e)}")


@router.post("/validate-component")
async def validate_component(req: ValidateComponentRequest, _=Depends(verify_azure_token)):
    raw = invoke_claude(
        f"""Is "{req.name}" a real, provisionable AWS infrastructure component that can be managed with Terraform?
Answer with JSON only: {{"valid": true, "reason": "one line"}} or {{"valid": false, "reason": "one line explaining what it is not"}}
Real examples: EC2, RDS, S3, VPC, EKS, Lambda, ECS, ElastiCache, CloudFront, Route53, ALB, API Gateway, DynamoDB, SQS, SNS, KMS, WAF, Redshift.
If it is a made-up word, random text, or not an AWS service, return valid=false.""",
        max_tokens=80,
        temperature=0,
    )
    try:
        result = json.loads(strip_fences(raw, "json"))
    except Exception:
        result = {"valid": True, "reason": ""}
    return result


@router.post("/extract-components")
async def extract_components(req: ExtractRequest, _=Depends(verify_azure_token)):
    """
    Use Claude to read the SAD and return a JSON list of infra components.
    """
    clean_doc = strip_html(req.document_content)[:10000]

    prompt = f"""You are a cloud infrastructure expert. Read the Solution Architecture Document below and extract the infrastructure components that are EXPLICITLY mentioned as needing to be provisioned.

Rules:
- ONLY include components clearly named or described in the document. Do NOT infer or add components that are not mentioned.
- Deduplicate: if the same AWS service is mentioned multiple times, include it once.
- Do NOT include generic concepts (e.g. "internet", "user", "browser") — only provisionable AWS resources.
- Each component must have:
  - "id": unique snake_case identifier (e.g. "vpc", "eks_cluster", "rds_postgres")
  - "name": human-readable name (e.g. "VPC", "EKS Cluster", "RDS PostgreSQL")
  - "description": one-line description of its role in the architecture
  - "category": exactly one of: networking, compute, storage, database, security, integration, monitoring

Document: {req.document_title}

--- DOCUMENT START ---
{clean_doc}
--- DOCUMENT END ---

Return ONLY a JSON array. No explanation. No markdown fences. No preamble.

Example:
[
  {{"id": "vpc", "name": "VPC", "description": "Main virtual private cloud with public and private subnets", "category": "networking"}},
  {{"id": "eks_cluster", "name": "EKS Cluster", "description": "Managed Kubernetes cluster for microservices", "category": "compute"}}
]"""

    try:
        raw = invoke_claude(prompt, max_tokens=3000, temperature=0)
        raw = strip_fences(raw, "json")
        components = json.loads(raw)
        return {"components": components}
    except json.JSONDecodeError:
        # Try to extract JSON array from response
        match = re.search(r'\[.*\]', raw, re.DOTALL)
        if match:
            components = json.loads(match.group())
            return {"components": components}
        raise


@router.post("/generate-stream")
async def generate_terraform_stream(req: GenerateRequest, _=Depends(verify_azure_token)):
    """
    Stream Terraform module generation for selected components.
    Yields newline-delimited JSON objects: {file, content, done}
    """
    project   = re.sub(r'[^a-z0-9_]', '_', req.project_name.lower())
    region    = req.aws_region or "us-east-1"
    env       = req.environment or "dev"
    components = req.selected_components
    clean_doc  = strip_html(req.document_content)[:6000]

    component_list = "\n".join(
        f"- {c.name} ({c.id}): {c.description}" for c in components
    )

    def generate():
        files: dict[str, str] = {}
        valid_components = []  # tracks only components that passed validation

        # ── 1. Generate each module (one file per call to avoid truncation) ──────
        for comp in components:
            context = (
                f"Project: {project}\nEnvironment: {env}\nAWS Region: {region}\n"
                f"Component: {comp.name} ({comp.id}) — {comp.description}\n\n"
                f"Architecture context:\n{clean_doc[:2000]}"
            )
            base_rules = (
                "- Use Terraform AWS provider ~> 5.0\n"
                "- Use snake_case for all resource names\n"
                "- Add tags: Project, Environment, ManagedBy=Terraform\n"
                "- Use variables for all configurable values\n"
                "- Apply production best practices: encryption at rest, security groups with least privilege, multi-AZ where applicable\n"
                "- Use realistic, sensible defaults (e.g. db.t3.micro for dev RDS, t3.medium for EC2, /16 CIDR for VPC)\n"
                "- Reference other modules via variables when components depend on each other (e.g. subnet_ids, vpc_id, security_group_ids)\n"
                "- Include lifecycle rules, backup settings, and monitoring where relevant to this resource type\n"
            )

            all_components_context = f"All components in this architecture:\n{component_list}\n\n"

            try:
                # ── Validate component is a real AWS resource ──
                validation_raw = invoke_claude(
                    f"""Is "{comp.name}" (id: "{comp.id}", description: "{comp.description}") a real, provisionable AWS infrastructure component that can be managed with Terraform?

Answer with JSON only: {{"valid": true/false, "reason": "one line explanation"}}
If it is a real AWS service or resource (e.g. EC2, RDS, S3, VPC, EKS, Lambda, etc.), return valid=true.
If it is a made-up name, random text, or not a real AWS resource, return valid=false.""",
                    max_tokens=100,
                    temperature=0,
                )
                try:
                    validation = json.loads(strip_fences(validation_raw, "json"))
                except Exception:
                    validation = {"valid": True}

                if not validation.get("valid", True):
                    yield json.dumps({"type": "error", "module": comp.id, "message": f'"{comp.name}" is not a recognized AWS resource: {validation.get("reason", "unknown component")}'}) + "\n"
                    continue

                valid_components.append(comp)

                # ── main.tf ──
                main_raw = invoke_claude(
                    f"""You are a senior AWS infrastructure engineer writing production-grade Terraform.

{context}
{all_components_context}
Task: Write the main.tf for the "{comp.name}" module.

Requirements:
{base_rules}
- Read the architecture context carefully. Configure this resource to match what is described (e.g. if the SAD mentions "private subnets", set multi_az=true for RDS; if it mentions "high availability", add redundancy).
- If this component connects to others in the list (e.g. RDS needs a VPC, EC2 needs a subnet), accept those as input variables — do NOT hardcode IDs.
- Include real resource blocks with all important arguments filled in — not just stubs.
- Add inline comments explaining non-obvious configuration choices.

Return ONLY raw HCL. No explanation. No markdown fences.""",
                    max_tokens=6000,
                )
                files[f"modules/{comp.id}/main.tf"] = strip_fences(main_raw)

                # ── variables.tf ──
                variables_raw = invoke_claude(
                    f"""You are a senior AWS infrastructure engineer writing production-grade Terraform.

{context}
{all_components_context}
Task: Write the variables.tf for the "{comp.name}" module.

Requirements:
- Declare every variable referenced in main.tf with: type, description, and a sensible default.
- For IDs passed from other modules (vpc_id, subnet_ids, security_group_ids), set default = "" or default = [] with a clear description.
- Use realistic defaults that match the environment "{env}" (e.g. smaller instance sizes for dev, larger for prod).
- Group related variables with comments.

Return ONLY raw HCL. No explanation. No markdown fences.""",
                    max_tokens=4000,
                )
                files[f"modules/{comp.id}/variables.tf"] = strip_fences(variables_raw)

                # ── outputs.tf ──
                outputs_raw = invoke_claude(
                    f"""You are a senior AWS infrastructure engineer writing production-grade Terraform.

{context}
Task: Write the outputs.tf for the "{comp.name}" module.

Requirements:
- Export all values that other modules will need: IDs, ARNs, endpoints, DNS names, security group IDs.
- Add a description to each output so callers know what they're getting.
- Only output values that actually exist on the resources in main.tf.

Return ONLY raw HCL. No explanation. No markdown fences.""",
                    max_tokens=2000,
                )
                files[f"modules/{comp.id}/outputs.tf"] = strip_fences(outputs_raw)

                yield json.dumps({"type": "module_done", "module": comp.id, "name": comp.name}) + "\n"
            except Exception as e:
                logger.error(f"Module generation failed for {comp.id}: {e}")
                yield json.dumps({"type": "error", "module": comp.id, "message": str(e)}) + "\n"

        # ── 2. Generate root main.tf via Claude (wires inter-module deps) ────────
        valid_component_list = "\n".join(
            f"- {c.name} ({c.id}): {c.description}" for c in valid_components
        )
        if not valid_components:
            files["main.tf"] = "# No valid AWS components were generated."
        else:
         root_main_raw = invoke_claude(
            f"""You are a senior AWS infrastructure engineer writing production-grade Terraform.

Project: {project} | Environment: {env} | Region: {region}

Modules in this architecture:
{valid_component_list}

Architecture context:
{clean_doc[:3000]}

Task: Write the ROOT main.tf that:
1. Declares terraform block with required_version ">= 1.5.0", AWS provider ~> 5.0, and an S3 backend block (commented out) named "{project}-terraform-state".
2. Declares provider "aws" with region = var.aws_region and default_tags (Project, Environment, ManagedBy=Terraform).
3. Calls each module with source = "./modules/<id>" and passes the correct variables.
4. IMPORTANTLY: wire inter-module dependencies — e.g. if RDS exists, pass vpc_id = module.vpc.vpc_id and subnet_ids = module.vpc.private_subnet_ids. If EC2 exists, pass subnet_id and security_group_ids from the VPC module. Use module outputs to connect them.
5. Add inline comments explaining the dependency wiring.

Return ONLY raw HCL. No explanation. No markdown fences.""",
            max_tokens=4000,
         )
         files["main.tf"] = strip_fences(root_main_raw)

        # ── 3. Generate root variables.tf ──────────────────────────────────────
        root_vars = f'''variable "project" {{
  description = "Project name used for resource naming and tagging"
  type        = string
  default     = "{project}"
}}

variable "environment" {{
  description = "Deployment environment (dev, staging, prod)"
  type        = string
  default     = "{env}"
}}

variable "aws_region" {{
  description = "AWS region to deploy resources"
  type        = string
  default     = "{region}"
}}
'''
        files["variables.tf"] = root_vars

        # ── 4. Generate root outputs.tf ────────────────────────────────────────
        output_blocks = ""
        for comp in components:
            output_blocks += f'''
output "{comp.id}_outputs" {{
  description = "Outputs from the {comp.name} module"
  value       = module.{comp.id}
  sensitive   = false
}}
'''
        files["outputs.tf"] = output_blocks.strip()

        # ── 5. terraform.tfvars.example ────────────────────────────────────────
        files["terraform.tfvars.example"] = f'''# Copy this file to terraform.tfvars and fill in your values

project     = "{project}"
environment = "{env}"
aws_region  = "{region}"
'''

        # ── 6. README.md ───────────────────────────────────────────────────────
        module_table = "\n".join(
            f"| `{c.id}` | {c.name} | {c.description} |"
            for c in components
        )
        files["README.md"] = f'''# {project.replace("_", " ").title()} — Terraform Infrastructure

Generated from: {req.document_content[:60].strip()}...

## Modules

| Module | Name | Description |
|--------|------|-------------|
{module_table}

## Usage

```bash
# Initialize
terraform init

# Review plan
terraform plan -var-file=terraform.tfvars

# Apply
terraform apply -var-file=terraform.tfvars
```

## Requirements

- Terraform >= 1.5.0
- AWS CLI configured
- AWS Provider ~> 5.0
'''

        # Stream all files at the end
        yield json.dumps({"type": "files", "files": files}) + "\n"
        yield json.dumps({"type": "done"}) + "\n"

    return StreamingResponse(generate(), media_type="application/x-ndjson")


@router.post("/download")
async def download_terraform_zip(req: DownloadRequest, _=Depends(verify_azure_token)):
    """
    Package all generated Terraform files into a ZIP and return for download.
    """
    project = re.sub(r'[^a-z0-9_]', '_', (req.project_name or "terraform").lower())
    buf = io.BytesIO()

    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for filename, content in req.files.items():
            zf.writestr(f"{project}/{filename}", content)

    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{project}-terraform.zip"'},
    )


@router.post("/push-github")
async def push_to_github(req: GitHubPushRequest, token_data: dict = Depends(verify_azure_token)):
    """
    Push all generated Terraform files to a GitHub repository in one commit.
    Uses the Git Trees API to batch all files in a single push.
    Creates the repo if it doesn't exist.
    """
    user = token_data.get("preferred_username") or token_data.get("upn") or token_data.get("sub", "unknown")

    # Accept full GitHub URLs — extract just the repo name
    # e.g. https://github.com/user/my-repo.git  →  my-repo
    repo_name = req.repo_name.strip()
    if "github.com" in repo_name:
        repo_name = repo_name.rstrip("/").removesuffix(".git").split("/")[-1]
    repo_name = repo_name.removesuffix(".git").strip()
    if not repo_name:
        raise HTTPException(status_code=400, detail="Invalid repository name.")

    logger.info(f"push-github called by {user}, repo={repo_name}, files={len(req.files)}")

    gh_headers = {
        "Authorization": f"Bearer {req.github_token}",
        "Accept": "application/vnd.github+json",
        "X-GitHub-Api-Version": "2022-11-28",
    }
    base_url = "https://api.github.com"

    async with httpx.AsyncClient(timeout=30, headers=gh_headers) as client:

        # ── 1. Get authenticated user ──────────────────────────────────────
        user_resp = await client.get(f"{base_url}/user")
        if user_resp.status_code != 200:
            logger.error(f"GitHub token validation failed: {user_resp.status_code} {user_resp.text[:200]}")
            raise HTTPException(status_code=400, detail="Invalid GitHub token. Make sure it has 'repo' scope and is not expired.")
        owner = user_resp.json()["login"]

        # ── 2. Create repo if it doesn't exist ────────────────────────────
        repo_resp = await client.get(f"{base_url}/repos/{owner}/{repo_name}")
        if repo_resp.status_code == 404:
            create_resp = await client.post(
                f"{base_url}/user/repos",
                json={
                    "name": repo_name,
                    "private": req.private,
                    "auto_init": True,
                    "description": "Terraform infrastructure code — generated by SDLC Orchestrator",
                },
            )
            if create_resp.status_code not in (200, 201):
                raise HTTPException(status_code=400, detail=f"Failed to create repo '{repo_name}': {create_resp.text[:300]}")
        elif repo_resp.status_code != 200:
            raise HTTPException(status_code=400, detail=f"GitHub error accessing repo: {repo_resp.text[:300]}")

        # ── 3. Get HEAD commit SHA of branch ──────────────────────────────
        ref_resp = await client.get(f"{base_url}/repos/{owner}/{repo_name}/git/ref/heads/{req.branch}")
        if ref_resp.status_code == 404:
            # Branch doesn't exist — get default branch HEAD
            for default_branch in ["main", "master"]:
                default_resp = await client.get(f"{base_url}/repos/{owner}/{repo_name}/git/ref/heads/{default_branch}")
                if default_resp.status_code == 200:
                    break
            else:
                raise HTTPException(status_code=400, detail="Could not find default branch. The repo may still be initializing — wait a moment and try again.")
            ref_data = default_resp.json()
            if "object" not in ref_data:
                raise HTTPException(status_code=400, detail=f"Unexpected GitHub response: {str(ref_data)[:200]}")
            base_sha = ref_data["object"]["sha"]
            # Create the new branch
            await client.post(
                f"{base_url}/repos/{owner}/{repo_name}/git/refs",
                json={"ref": f"refs/heads/{req.branch}", "sha": base_sha},
            )
        elif ref_resp.status_code == 200:
            ref_data = ref_resp.json()
            if "object" not in ref_data:
                raise HTTPException(status_code=400, detail=f"Unexpected GitHub ref response: {str(ref_data)[:200]}")
            base_sha = ref_data["object"]["sha"]
        elif ref_resp.status_code == 409:
            # 409 = repo exists but is completely empty (no commits yet)
            base_sha = None
        else:
            raise HTTPException(status_code=400, detail=f"Failed to get branch ref: {ref_resp.text[:300]}")

        is_empty_repo = base_sha is None

        if is_empty_repo:
            # ── Empty repo: Git Data API (blobs/trees) doesn't work yet.
            # Use Contents API instead — creates one commit per file but
            # correctly initialises the repo.
            prefix = (req.folder_prefix or "").strip("/")
            commit_sha = None
            for filename, content in req.files.items():
                path = f"{prefix}/{filename}" if prefix else filename
                put_payload: dict = {
                    "message": req.commit_message if commit_sha is None else f"chore: add {filename}",
                    "content": base64.b64encode(content.encode()).decode(),
                    "branch": req.branch,
                }
                put_resp = await client.put(
                    f"{base_url}/repos/{owner}/{repo_name}/contents/{path}",
                    json=put_payload,
                )
                if put_resp.status_code not in (200, 201):
                    raise HTTPException(status_code=400, detail=f"Failed to push {filename}: {put_resp.text[:200]}")
                commit_sha = put_resp.json()["commit"]["sha"]

        else:
            # ── Existing repo: batch all files into one commit via Git Trees API
            prefix = (req.folder_prefix or "").strip("/")
            tree_items = []
            for filename, content in req.files.items():
                path = f"{prefix}/{filename}" if prefix else filename
                blob_resp = await client.post(
                    f"{base_url}/repos/{owner}/{repo_name}/git/blobs",
                    json={
                        "content": base64.b64encode(content.encode()).decode(),
                        "encoding": "base64",
                    },
                )
                if blob_resp.status_code not in (200, 201):
                    raise HTTPException(status_code=400, detail=f"Failed to create blob for {filename}: {blob_resp.text[:200]}")
                tree_items.append({
                    "path": path,
                    "mode": "100644",
                    "type": "blob",
                    "sha": blob_resp.json()["sha"],
                })

            # Create tree
            tree_resp = await client.post(
                f"{base_url}/repos/{owner}/{repo_name}/git/trees",
                json={"base_tree": base_sha, "tree": tree_items},
            )
            if tree_resp.status_code not in (200, 201):
                raise HTTPException(status_code=400, detail=f"Failed to create tree: {tree_resp.text[:300]}")
            tree_sha = tree_resp.json()["sha"]

            # Create commit
            commit_resp = await client.post(
                f"{base_url}/repos/{owner}/{repo_name}/git/commits",
                json={"message": req.commit_message, "tree": tree_sha, "parents": [base_sha]},
            )
            if commit_resp.status_code not in (200, 201):
                raise HTTPException(status_code=400, detail=f"Failed to create commit: {commit_resp.text[:300]}")
            commit_sha = commit_resp.json()["sha"]

            # Update branch ref
            update_resp = await client.patch(
                f"{base_url}/repos/{owner}/{repo_name}/git/refs/heads/{req.branch}",
                json={"sha": commit_sha, "force": False},
            )
            if update_resp.status_code not in (200, 201):
                raise HTTPException(status_code=400, detail=f"Failed to update branch: {update_resp.text[:300]}")

        repo_url = f"https://github.com/{owner}/{repo_name}"
        return {
            "success": True,
            "repo_url": repo_url,
            "branch": req.branch,
            "commit_sha": commit_sha[:7],
            "files_pushed": len(req.files),
            "message": f"Successfully pushed {len(req.files)} files to {repo_url}/tree/{req.branch}",
        }


@router.post("/push-harness")
async def push_to_harness(req: HarnessPushRequest, token_data: dict = Depends(verify_azure_token)):
    """
    Push all generated Terraform files to a Harness Code Repository in one commit.
    Uses the Harness Code commits API (Gitness-based) to batch all files.
    Handles both new files (CREATE) and existing files (UPDATE with SHA).
    """
    user = token_data.get("preferred_username") or token_data.get("upn") or token_data.get("sub", "unknown")

    # Parse repo_ref from the full Harness Code URL.
    # Supported URL formats:
    #   https://app.harness.io/ng/account/{accountId}/module/code/orgs/{orgId}/projects/{projectId}/repos/{repoName}
    #   https://app.harness.io/code/{accountId}/orgs/{orgId}/projects/{projectId}/repos/{repoName}
    #   plain:  {accountId}/{orgId}/{projectId}/{repoName}
    url = req.repo_url.strip().rstrip("/")
    m = re.search(
        r'account[s]?/([^/]+)/(?:module/code/)?orgs/([^/]+)/projects/([^/]+)/repos/([^/?#]+)',
        url,
    )
    if m:
        account_id, org_id, project_id, repo_name = m.group(1), m.group(2), m.group(3), m.group(4)
    else:
        # Try plain slash-separated format
        parts = [p for p in url.split("/") if p]
        if len(parts) == 4:
            account_id, org_id, project_id, repo_name = parts
        else:
            raise HTTPException(
                status_code=400,
                detail=(
                    "Could not parse Harness repo URL. Paste the full URL from your browser, e.g. "
                    "https://app.harness.io/ng/account/ACCOUNT/module/code/orgs/ORG/projects/PROJECT/repos/REPO"
                ),
            )

    repo_ref = f"{account_id}/{org_id}/{project_id}/{repo_name}"
    base_url = "https://app.harness.io/code/api/v1"
    headers = {
        "x-api-key": req.api_key,
        "Content-Type": "application/json",
    }
    prefix = (req.folder_prefix or "").strip("/")

    logger.info(f"push-harness called by {user}, repo={repo_ref}, files={len(req.files)}")

    async with httpx.AsyncClient(timeout=60, headers=headers) as client:

        # ── 1. Check repo — create automatically if it doesn't exist ─────────
        repo_resp = await client.get(f"{base_url}/repos/{repo_ref}")
        if repo_resp.status_code == 401:
            raise HTTPException(status_code=400, detail="Invalid Harness API key or insufficient permissions.")
        if repo_resp.status_code == 404:
            # Repo doesn't exist — create it
            logger.info(f"Harness repo {repo_ref} not found, creating it")
            parent_ref = f"{account_id}/{org_id}/{project_id}"
            create_resp = await client.post(
                f"{base_url}/repos",
                json={
                    "identifier": repo_name,
                    "parent_ref": parent_ref,
                    "default_branch": req.branch or "main",
                    "description": "Terraform infrastructure code — generated by SDLC Orchestrator",
                    "is_public": False,
                },
            )
            if create_resp.status_code not in (200, 201):
                raise HTTPException(
                    status_code=400,
                    detail=f"Repository not found and auto-creation failed: {create_resp.text[:300]}",
                )
            logger.info(f"Created Harness repo {repo_ref}")
        elif repo_resp.status_code != 200:
            raise HTTPException(status_code=400, detail=f"Failed to access Harness repo: {repo_resp.text[:300]}")

        # ── 2. Build CREATE actions for all files ─────────────────────────────
        actions = []
        for filename, content in req.files.items():
            path = f"{prefix}/{filename}" if prefix else filename
            actions.append({
                "action": "CREATE",
                "path": path,
                "payload": base64.b64encode(content.encode()).decode(),
                "encoding": "base64",
            })

        commit_body = {
            "branch": req.branch,
            "new_branch": "",
            "title": req.commit_message,
            "message": f"Generated by SDLC Orchestrator — {len(req.files)} Terraform files",
            "actions": actions,
        }

        # ── 3. Attempt batch commit ───────────────────────────────────────────
        commit_resp = await client.post(f"{base_url}/repos/{repo_ref}/commits", json=commit_body)

        if commit_resp.status_code == 201:
            data = commit_resp.json()
            commit_sha = data.get("sha", "")
            repo_url = f"https://app.harness.io/code/{repo_ref}/commits"
            return {
                "success": True,
                "repo_url": repo_url,
                "branch": req.branch,
                "commit_sha": commit_sha[:7] if commit_sha else "",
                "files_pushed": len(req.files),
                "message": f"Successfully pushed {len(req.files)} files to Harness Code",
            }

        # ── 4. Some files already exist — get their SHAs and retry with UPDATE ─
        if commit_resp.status_code in (409, 422):
            logger.info(f"Some files exist in Harness repo, retrying with UPDATE actions")
            update_actions = []
            for action in actions:
                file_path = action["path"]
                file_resp = await client.get(
                    f"{base_url}/repos/{repo_ref}/content/{file_path}",
                    params={"ref": req.branch},
                )
                if file_resp.status_code == 200:
                    sha = file_resp.json().get("sha", "")
                    update_actions.append({**action, "action": "UPDATE", "sha": sha})
                else:
                    update_actions.append(action)  # Keep as CREATE if not found

            retry_body = {**commit_body, "actions": update_actions}
            retry_resp = await client.post(f"{base_url}/repos/{repo_ref}/commits", json=retry_body)
            if retry_resp.status_code == 201:
                data = retry_resp.json()
                commit_sha = data.get("sha", "")
                repo_url = f"https://app.harness.io/code/{repo_ref}/commits"
                return {
                    "success": True,
                    "repo_url": repo_url,
                    "branch": req.branch,
                    "commit_sha": commit_sha[:7] if commit_sha else "",
                    "files_pushed": len(req.files),
                    "message": f"Successfully updated {len(req.files)} files in Harness Code",
                }
            raise HTTPException(status_code=400, detail=f"Harness commit failed: {retry_resp.text[:400]}")

        raise HTTPException(status_code=400, detail=f"Harness commit failed ({commit_resp.status_code}): {commit_resp.text[:400]}")
