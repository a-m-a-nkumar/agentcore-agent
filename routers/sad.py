"""
FastAPI router for the SAD phase of the multi-session Design Assistant.

All endpoints below are scoped by `session_id` (from `design_sessions`).
Heavy work (intent routing, generation, audit, edit) is delegated to the
`sdlc-dev-sad-orchestrator` Lambda via boto3. Light reads (sections,
diagrams, facts, DOCX export) stay in this router and read directly from
S3.

Endpoints (all under /api/sad):
  POST    /turn                      â†’ unified chat box (multipart, JSON or SSE-style chunks)
  POST    /generate                  â†’ kick off section workers (returns final SAD JSON)
  POST    /audit                     â†’ run audit, persist results, return badges + details
  POST    /revert-section            â†’ pop one entry off a section's previous_versions stack
  GET     /{session_id}/sections     â†’ first paint
  GET     /{session_id}/section/{n}  â†’ single-section refresh
  GET     /{session_id}/diagram/{kind} â†’ returns SVG bytes (kind âˆˆ logical|security|infrastructure; v1 reuses logical)
  GET     /{session_id}/facts        â†’ facts panel
  GET     /download-sad/{session_id} â†’ DOCX export (cairosvg â†’ PNG, python-docx)
"""

import io
import json
import logging
import os
import re
import time
from typing import Any, Dict, List, Optional, Tuple

import boto3
from botocore.config import Config as BotoConfig
from docx import Document
from docx.shared import Inches
from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from fastapi.responses import Response, StreamingResponse
from pydantic import BaseModel

from db_helper import (
    get_design_session,
    get_diagram_slots,
    get_user_atlassian_credentials,
    track_event,
    update_design_session,
)
from environment import S3_BUCKET_NAME
from services.confluence_service import ConfluenceService
from services.s3_service import s3_put_object

# Reuse the projects-router auth dependency (DB user row with key "id")
from .projects import get_current_user

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/sad", tags=["sad"])

AWS_REGION = os.getenv("AWS_REGION", "us-east-1")
SAD_LAMBDA_NAME = os.getenv("LAMBDA_SAD_ORCHESTRATOR", "sdlc-dev-sad-orchestrator")

_lambda_client = None
_s3_client = None


def _lambda():
    global _lambda_client
    if _lambda_client is None:
        _lambda_client = boto3.client(
            "lambda", region_name=AWS_REGION,
            config=BotoConfig(read_timeout=300, connect_timeout=20, retries={"max_attempts": 1}),
        )
    return _lambda_client


def _s3():
    global _s3_client
    if _s3_client is None:
        _s3_client = boto3.client("s3", region_name=AWS_REGION)
    return _s3_client


def _ensure_session_owned(session_id: str, user_id: str) -> Dict[str, Any]:
    s = get_design_session(session_id)
    if not s:
        raise HTTPException(status_code=404, detail="Session not found")
    if s["user_id"] != user_id:
        raise HTTPException(status_code=403, detail="You don't have access to this session")
    return s


def _invoke_sad_lambda(payload: Dict[str, Any]) -> Dict[str, Any]:
    resp = _lambda().invoke(
        FunctionName=SAD_LAMBDA_NAME,
        InvocationType="RequestResponse",
        Payload=json.dumps(payload).encode("utf-8"),
    )
    body_bytes = resp["Payload"].read()
    if "FunctionError" in resp:
        raise HTTPException(
            status_code=502,
            detail=f"SAD Lambda error: {body_bytes.decode('utf-8', errors='replace')[:500]}",
        )
    try:
        outer = json.loads(body_bytes)
        if "body" in outer and isinstance(outer["body"], str):
            return json.loads(outer["body"])
        return outer
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"SAD Lambda response parse error: {e}")


# ============================================
# Confluence URL ingestion helpers
# ============================================
#
# When a user pastes a Confluence URL into the SAD chat, we treat it the
# same way as an uploaded file: fetch the page body, strip the HTML, build
# a `{filename, extracted_text}` payload, and forward it to the Lambda's
# INGEST_DOC handler. Multiple URLs in one message are supported â€” they
# become multiple file_payloads, the Lambda emits one doc_ingested card
# per URL, and the auto-regen flag fires only on the last card.

# Atlassian Cloud only â€” `https://<tenant>.atlassian.net/wiki/spaces/<SPACE>/pages/<PAGE_ID>(/...)`.
# Confluence Server / DC out of scope per plan decision.
_CONFLUENCE_URL_PATTERN = re.compile(
    r"https?://([\w\-]+\.atlassian\.net)/wiki/spaces/[^/\s]+/pages/(\d+)(?:/[^\s]*)?",
    re.IGNORECASE,
)


def _extract_confluence_urls(message: str) -> List[Tuple[str, str]]:
    """Return [(domain, page_id), ...] in the order they appear; deduped by page_id."""
    if not message:
        return []
    seen: set[str] = set()
    out: List[Tuple[str, str]] = []
    for m in _CONFLUENCE_URL_PATTERN.finditer(message):
        domain = m.group(1).lower()
        page_id = m.group(2)
        if page_id in seen:
            continue
        seen.add(page_id)
        out.append((domain, page_id))
    return out


def _strip_confluence_html(html: str) -> str:
    """Strip Confluence storage HTML to plain text.

    Reuses the same approach as `services.rag_service._strip_html` â€”
    drop tags, collapse whitespace. Inlined here to avoid pulling the
    full RAG service into this module's import graph.
    """
    if not html:
        return ""
    text = re.sub(r"<[^>]+>", " ", html)
    return re.sub(r"\s+", " ", text).strip()


class _ConfluenceFetchResult:
    """Lightweight tagged-union for what `_fetch_confluence_page` returns."""

    __slots__ = ("status", "file_payload", "page_title", "domain", "reason")

    def __init__(
        self,
        status: str,
        *,
        file_payload: Optional[Dict[str, Any]] = None,
        page_title: Optional[str] = None,
        domain: Optional[str] = None,
        reason: Optional[str] = None,
    ):
        self.status = status  # "ok" | "tenant_mismatch" | "not_linked" | "fetch_failed" | "empty"
        self.file_payload = file_payload
        self.page_title = page_title
        self.domain = domain
        self.reason = reason


def _fetch_confluence_page(
    domain: str,
    page_id: str,
    user_id: str,
    *,
    session_id: Optional[str] = None,
) -> _ConfluenceFetchResult:
    """Fetch one Confluence page, return a tagged result.

    On `status="ok"`, `file_payload` is shaped like the file-upload
    payload (`{filename, extracted_text}`) so it can be passed to the
    Lambda's INGEST_DOC path with no further translation.
    """
    creds = get_user_atlassian_credentials(user_id)
    if not creds or not creds.get("atlassian_api_token"):
        return _ConfluenceFetchResult("not_linked", domain=domain)

    linked_domain = (creds.get("atlassian_domain") or "").lower()
    if linked_domain and linked_domain != domain:
        logger.info(f"[SAD] confluence tenant mismatch: requested={domain}, linked={linked_domain}")
        return _ConfluenceFetchResult("tenant_mismatch", domain=domain)

    try:
        confluence = ConfluenceService(
            domain,
            creds["atlassian_email"],
            creds["atlassian_api_token"],
        )
        page = confluence.get_content_page_by_id(page_id, expand="body.storage,version")
    except Exception as e:
        logger.warning(f"[SAD] confluence fetch failed: page_id={page_id} domain={domain} err={e}")
        return _ConfluenceFetchResult("fetch_failed", domain=domain, reason=str(e))

    title = (page.get("title") or "Confluence page").strip()
    body_html = ((page.get("body") or {}).get("storage") or {}).get("value") or ""
    plain = _strip_confluence_html(body_html)
    if not plain:
        return _ConfluenceFetchResult("empty", domain=domain, page_title=title)

    # Audit trail: persist to the same `sources/` prefix file uploads use.
    if session_id:
        try:
            safe_title = re.sub(r"[^\w\-. ]+", "_", title)[:120].strip() or page_id
            key = f"sessions/{session_id}/sources/{safe_title}.confluence.txt"
            s3_put_object(key=key, body=plain.encode("utf-8"), content_type="text/plain")
        except Exception as e:
            logger.warning(f"[SAD] failed to persist confluence source to S3 (non-fatal): {e}")

    filename = f"{title}.confluence"
    logger.info(
        f"[SAD] confluence ingest: domain={domain} page_id={page_id} title={title!r} chars={len(plain)}"
    )
    return _ConfluenceFetchResult(
        "ok",
        file_payload={
            "filename": filename,
            "extracted_text": plain,
            "doc_id": f"confluence:{page_id}",
        },
        page_title=title,
        domain=domain,
    )


def _confluence_warning_card(reason: str, *, domain: Optional[str] = None, page_id: Optional[str] = None) -> Dict[str, Any]:
    """Synthesize a `text` card the frontend renders as a chat bubble.

    Covers the cases where we couldn't ingest a URL: not linked, wrong
    tenant, fetch failed, or empty page. Returned to the frontend as
    part of the `cards` array prepended to whatever the Lambda emits.
    """
    text_by_reason = {
        "not_linked": "To ingest Confluence pages from chat, link your Atlassian account in Settings.",
        "tenant_mismatch": (
            f"That URL is in another Atlassian tenant ({domain}) â€” "
            "link your account there to ingest from it. The rest of your message was processed normally."
        ),
        "fetch_failed": (
            f"Couldn't fetch Confluence page {page_id} (it may be deleted or you may not have access). "
            "The rest of your message was processed normally."
        ),
        "empty": (
            f"Confluence page {page_id} has no body content to ingest. "
            "The rest of your message was processed normally."
        ),
    }
    return {
        "type": "text",
        "payload": {"text": text_by_reason.get(reason, "Couldn't ingest a Confluence URL.")},
    }


# ============================================
# Models
# ============================================

class SADGenerateRequest(BaseModel):
    session_id: str
    project_id: Optional[str] = None
    brd_id: Optional[str] = None


class SADAuditRequest(BaseModel):
    session_id: str
    project_id: Optional[str] = None
    section_number: Optional[int] = None  # If set, audit only that section


class SADRevertRequest(BaseModel):
    session_id: str
    section_number: int


# ============================================
# Endpoints
# ============================================

@router.post("/turn")
async def sad_turn(
    session_id: str = Form(...),
    message: str = Form(""),
    project_id: Optional[str] = Form(None),
    viewing_section: Optional[int] = Form(None),
    last_card_type: Optional[str] = Form(None),
    last_proposed_section: Optional[int] = Form(None),
    file: Optional[UploadFile] = File(None),
    current_user: dict = Depends(get_current_user),
):
    """One chat-box turn. The Lambda always sees a single source of input:
      â€¢ An uploaded file â†’ INGEST_DOC (legacy single-file path).
      â€¢ One or more Confluence URLs in the message â†’ INGEST_DOC for each
        (multi-file path; Lambda iterates and emits one card per page).
      â€¢ Neither â†’ plain-text turn (intent router classifies).

    File upload takes priority â€” if a file is attached, URLs in the
    message are ignored that turn. Response shape is always
    `{cards: [...]}` so the frontend can render multiple bubbles per turn.
    """
    user_id = current_user["id"]
    session = _ensure_session_owned(session_id, user_id)

    # â”€â”€ 1. Single file upload (legacy path) â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    file_payload: Optional[Dict[str, Any]] = None
    if file is not None:
        try:
            from app import extract_text  # reuse the project's extractor
        except Exception:
            extract_text = None
        raw = await file.read()
        text = ""
        if extract_text:
            try:
                text = extract_text(raw, file.filename or "uploaded")
            except Exception as e:
                logger.warning(f"[SAD] extract_text failed for {file.filename}: {e}")
        file_payload = {
            "filename": file.filename or "uploaded",
            "extracted_text": text,
        }
        try:
            key = f"sessions/{session_id}/sources/{file.filename or 'uploaded'}"
            s3_put_object(key=key, body=raw, content_type=file.content_type or "application/octet-stream")
            file_payload["s3_key"] = key
        except Exception as e:
            logger.warning(f"[SAD] failed to persist source file {file.filename}: {e}")

    # â”€â”€ 2. Confluence URLs in the message (skipped when a file is attached) â”€â”€
    files_payload: List[Dict[str, Any]] = []
    pre_cards: List[Dict[str, Any]] = []
    if file_payload is None and message:
        url_matches = _extract_confluence_urls(message)
        if url_matches:
            for domain, page_id in url_matches:
                result = _fetch_confluence_page(domain, page_id, user_id, session_id=session_id)
                if result.status == "ok" and result.file_payload:
                    files_payload.append(result.file_payload)
                else:
                    pre_cards.append(
                        _confluence_warning_card(result.status, domain=domain, page_id=page_id)
                    )

    # â”€â”€ 3. Forward to Lambda â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    payload = {
        "action": "turn",
        "session_id": session_id,
        "project_id": project_id or session.get("project_id"),
        "user_id": user_id,
        "message": message,
        "viewing_section": viewing_section,
        "last_card_type": last_card_type,
        "last_proposed_section": last_proposed_section,
        "stage": session.get("stage"),
        "file": file_payload,
        "files": files_payload or None,
    }
    lambda_result = _invoke_sad_lambda(payload)

    # â”€â”€ 4. Normalise response into {cards: [...]} â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    # Lambda may return either:
    #   â€¢ Legacy single card: {"type": "...", "payload": {...}}
    #   â€¢ Multi-ingest envelope: {"cards": [{"type", "payload"}, ...]}
    if isinstance(lambda_result, dict) and isinstance(lambda_result.get("cards"), list):
        lambda_cards = lambda_result["cards"]
    else:
        lambda_cards = [lambda_result] if isinstance(lambda_result, dict) else []
    all_cards = pre_cards + lambda_cards

    # â”€â”€ 5. Stage transitions based on the most-significant card â”€â”€â”€â”€â”€â”€â”€â”€
    # Use the LAST card so a doc_ingested with auto_regen=true correctly
    # transitions to SAD_GATHERING / SAD_GENERATING.
    last_card_type_resp = (all_cards[-1] or {}).get("type") if all_cards else None
    if last_card_type_resp == "generation_starting":
        update_design_session(session_id=session_id, stage="SAD_GENERATING", bump_activity=True)
    elif session.get("stage") in ("NEW", "DIAGRAM_READY") and last_card_type_resp in ("fact_saved", "doc_ingested"):
        update_design_session(session_id=session_id, stage="SAD_GATHERING", bump_activity=True)
    else:
        update_design_session(session_id=session_id, bump_activity=True)

    return {"cards": all_cards}


@router.post("/generate")
def sad_generate(
    body: SADGenerateRequest,
    current_user: dict = Depends(get_current_user),
):
    """Kick off the section workers. Returns the final SAD JSON when done.

    Client-side, we recommend showing a loading overlay; the call may take
    60-120 seconds. A streaming variant can be added later if SSE becomes
    necessary.
    """
    session = _ensure_session_owned(body.session_id, current_user["id"])
    update_design_session(session_id=body.session_id, stage="SAD_GENERATING", bump_activity=True)

    # SAD-redesign: surface per-type diagram slots to the Lambda so each
    # section reads from the matching slot (logicalâ†’Â§4, securityâ†’Â§6,
    # infrastructureâ†’Â§7). Missing or skipped slots render an explicit
    # placeholder, never silently substitute another type's artifact.
    try:
        slots_state = get_diagram_slots(body.session_id)
    except Exception as e:
        logger.warning(f"[SAD] failed to load diagram slots for {body.session_id}: {e}")
        slots_state = {"tool": None, "slots": {}}

    payload = {
        "action": "generate_sad",
        "session_id": body.session_id,
        "project_id": body.project_id or session.get("project_id"),
        "brd_id": body.brd_id,
        "user_id": current_user["id"],
        "diagram_slots": slots_state.get("slots") or {},
    }
    result = _invoke_sad_lambda(payload)
    update_design_session(session_id=body.session_id, stage="SAD_REFINING", bump_activity=True)
    try:
        track_event(
            current_user["id"],
            module="architecture",
            event_type="sad_generated",
            project_id=body.project_id or session.get("project_id"),
            metadata={
                "session_id": body.session_id,
                "brd_id": body.brd_id,
                "slots_present": list((slots_state.get("slots") or {}).keys()),
            },
        )
    except Exception as _track_err:
        logger.warning(f"track_event failed (non-fatal): {_track_err}")
    return result


@router.post("/audit")
def sad_audit(
    body: SADAuditRequest,
    current_user: dict = Depends(get_current_user),
):
    """Run audit on demand. When `section_number` is set, only that section
    is audited (cheap â€” one LLM call). Otherwise the full 10-section audit
    runs in parallel. Lambda persists audit_latest.json + decorates the
    section objects in sad_structure.json."""
    session = _ensure_session_owned(body.session_id, current_user["id"])
    payload = {
        "action": "audit",
        "session_id": body.session_id,
        "project_id": body.project_id or session.get("project_id"),
        "user_id": current_user["id"],
        "section_number": body.section_number,
    }
    return _invoke_sad_lambda(payload)


@router.post("/revert-section")
def sad_revert(
    body: SADRevertRequest,
    current_user: dict = Depends(get_current_user),
):
    _ensure_session_owned(body.session_id, current_user["id"])
    payload = {
        "action": "revert_section",
        "session_id": body.session_id,
        "section_number": body.section_number,
        "user_id": current_user["id"],
    }
    return _invoke_sad_lambda(payload)


# ============================================
# Manual section edit (no LLM, direct content overwrite)
# ============================================

class SADSaveSectionRequest(BaseModel):
    session_id: str
    section_number: int
    content: List[Dict[str, Any]]  # array of content blocks


@router.post("/save-section")
def sad_save_section(
    body: SADSaveSectionRequest,
    current_user: dict = Depends(get_current_user),
):
    """Persist user-edited section content directly. Bypasses the LLM â€”
    the user has typed the exact content they want. Pushes the previous
    content onto the section's `previous_versions` stack so the Revert
    button still works to undo this edit."""
    _ensure_session_owned(body.session_id, current_user["id"])
    sad = _read_sad(body.session_id)
    sections = sad.get("sections", [])
    if not (1 <= body.section_number <= len(sections)):
        raise HTTPException(status_code=404, detail=f"Section {body.section_number} not found")

    # Validate content shape â€” each block must have a recognised `type`.
    valid_types = {"paragraph", "heading", "ordered_list", "bullet_list", "table", "diagram"}
    for i, block in enumerate(body.content):
        if not isinstance(block, dict) or block.get("type") not in valid_types:
            raise HTTPException(
                status_code=400,
                detail=f"content[{i}].type must be one of {sorted(valid_types)}",
            )

    section = sections[body.section_number - 1]
    prev_content = section.get("content") or []

    # No-op detection: if the user clicked Save without changing anything,
    # don't push a duplicate onto the version stack and don't bump status
    # / last_modified_ts. Returning the section unchanged keeps Revert
    # meaningful (one undo == one real edit).
    unchanged = json.dumps(prev_content, sort_keys=True) == json.dumps(body.content, sort_keys=True)
    if unchanged:
        logger.info(
            f"[SAD] section {body.section_number} save: no-op (content identical), skipping version push"
        )
        return section

    if prev_content:
        stack = section.setdefault("previous_versions", [])
        stack.insert(0, prev_content)
        section["previous_versions"] = stack[:5]
    section["content"] = body.content
    section["status"] = "user_edited"
    section["last_modified_ts"] = int(time.time())

    # Write back to S3 via the centralized helper â€” the bucket policy
    # explicit-denies any PutObject that isn't SSE-KMS with our key.
    try:
        s3_put_object(
            key=f"sessions/{body.session_id}/sad/sad_structure.json",
            body=json.dumps(sad).encode("utf-8"),
            content_type="application/json",
        )
    except Exception as e:
        logger.error(f"[SAD] save-section S3 write failed: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to persist section: {e}")

    update_design_session(session_id=body.session_id, bump_activity=True)
    logger.info(
        f"[SAD] section {body.section_number} manually edited "
        f"(session {body.session_id}, {len(body.content)} blocks)"
    )
    return section


# ============================================
# Reads â€” direct S3 (no Lambda)
# ============================================

def _read_sad(session_id: str) -> Dict[str, Any]:
    try:
        obj = _s3().get_object(
            Bucket=S3_BUCKET_NAME,
            Key=f"sessions/{session_id}/sad/sad_structure.json",
        )
        return json.loads(obj["Body"].read().decode("utf-8"))
    except Exception:
        raise HTTPException(status_code=404, detail="SAD has not been generated for this session")


@router.get("/{session_id}/sections")
def get_sections(
    session_id: str,
    current_user: dict = Depends(get_current_user),
):
    _ensure_session_owned(session_id, current_user["id"])
    sad = _read_sad(session_id)
    return {
        "sad_id": sad.get("sad_id"),
        "stage": sad.get("stage"),
        "sections": [
            {
                "number": s["number"],
                "title": s["title"],
                "status": s.get("status"),
                "audit": s.get("audit"),
                "last_modified_ts": s.get("last_modified_ts"),
            }
            for s in sad.get("sections", [])
        ],
    }


@router.get("/{session_id}/section/{n}")
def get_section(
    session_id: str,
    n: int,
    current_user: dict = Depends(get_current_user),
):
    _ensure_session_owned(session_id, current_user["id"])
    sad = _read_sad(session_id)
    secs = sad.get("sections", [])
    if not (1 <= n <= len(secs)):
        raise HTTPException(status_code=404, detail=f"Section {n} not found")
    return secs[n - 1]


@router.get("/{session_id}/diagram/{kind}")
def get_diagram(
    session_id: str,
    kind: str,
    current_user: dict = Depends(get_current_user),
):
    """Serve the per-type SVG for a SAD diagram-block.

    P3 (honour every saved diagram in its own section): each type reads its
    own slot. Â§4 â†’ logical.svg, Â§6 â†’ security.svg, Â§7 â†’ infrastructure.svg.
    Falls back to logical.svg ONLY when explicitly asked for `logical` and the
    per-type file is missing â€” for the other two we 404 honestly so the SAD
    viewer can render an explicit "<view> not authored" placeholder rather
    than silently substituting the Logical artifact.
    """
    _ensure_session_owned(session_id, current_user["id"])
    if kind not in ("logical", "security", "infrastructure"):
        raise HTTPException(status_code=400, detail="kind must be logical | security | infrastructure")

    # PNG first (Lucid import + drawio rasterized export), SVG fallback
    # (drawio embedded export). Sniff in order; serve with the right MIME so
    # the browser doesn't guess.
    candidates = [
        (f"sessions/{session_id}/diagram/{kind}.png", "image/png"),
        (f"sessions/{session_id}/diagram/{kind}.svg", "image/svg+xml"),
    ]
    for primary_key, content_type in candidates:
        try:
            obj = _s3().get_object(Bucket=S3_BUCKET_NAME, Key=primary_key)
            return Response(content=obj["Body"].read(), media_type=content_type)
        except Exception as e:
            logger.debug(f"[SAD] diagram miss for {primary_key}: {e!r}")
            continue

    logger.info(f"[SAD] no diagram artifact for kind={kind} session={session_id}")
    # 404 â€” don't substitute another type's image. The SAD generator's
    # placeholder paragraph already covers skipped slots in the section
    # JSON; the viewer should render that placeholder instead of a
    # misleading image.
    raise HTTPException(
        status_code=404,
        detail=f"No {kind} diagram saved for this session",
    )


@router.get("/{session_id}/facts")
def get_facts(
    session_id: str,
    current_user: dict = Depends(get_current_user),
):
    _ensure_session_owned(session_id, current_user["id"])
    try:
        obj = _s3().get_object(
            Bucket=S3_BUCKET_NAME,
            Key=f"sessions/{session_id}/sad/facts.json",
        )
        return json.loads(obj["Body"].read().decode("utf-8"))
    except Exception:
        return {"sad_id": session_id, "facts": []}


# ============================================
# DOCX export
# ============================================

@router.get("/download-sad/{session_id}")
def download_sad(
    session_id: str,
    current_user: dict = Depends(get_current_user),
):
    """Render sad_structure.json into a DOCX. Diagram blocks are SVG â†’ PNG via cairosvg."""
    _ensure_session_owned(session_id, current_user["id"])
    sad = _read_sad(session_id)

    try:
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"python-docx not installed: {e}")

    try:
        import cairosvg
    except Exception:
        cairosvg = None  # We'll skip diagram rendering if it's missing.

    doc = Document()
    doc.add_heading("Software Architecture Document", level=0)
    doc.add_paragraph(f"Session ID: {session_id}")

    for section in sad.get("sections", []):
        doc.add_heading(f"{section['number']}. {section['title']}", level=1)
        for block in section.get("content", []) or []:
            t = block.get("type")
            if t == "paragraph":
                doc.add_paragraph(block.get("text", ""))
            elif t == "heading":
                doc.add_heading(block.get("text", ""), level=min(int(block.get("level", 3)), 4))
            elif t == "ordered_list":
                for it in block.get("items", []):
                    doc.add_paragraph(it, style="List Number")
            elif t == "bullet_list":
                for it in block.get("items", []):
                    doc.add_paragraph(it, style="List Bullet")
            elif t == "table":
                headers = block.get("headers", [])
                rows = block.get("rows", [])
                if headers:
                    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                    table.style = "Light Grid Accent 1"
                    for i, h in enumerate(headers):
                        table.cell(0, i).text = str(h)
                    for r_idx, r in enumerate(rows, start=1):
                        for c_idx, cell in enumerate(r[: len(headers)]):
                            table.cell(r_idx, c_idx).text = str(cell)
            elif t == "diagram":
                # Embed the per-section diagram. The section JSON's `s3_key`
                # already points to the correct file for THIS section (set
                # by lambda_sad_orchestrator._diagram_block_for_section from
                # the diagram_slots[type].artifact_key field), so we trust
                # that path and dispatch on the file extension:
                #
                #   - .png / .jpg / .jpeg  â†’  embed directly via python-docx
                #                             (Lucid imports + future drawio
                #                              PNG exports both land here)
                #   - .svg                 â†’  cairosvg â†’ PNG â†’ embed
                #                             (current drawio path)
                #
                # Earlier this code hardcoded `sessions/{id}/diagram/logical.png`
                # which silently fed the Â§4 logical image into Â§6 (security)
                # and Â§7 (infrastructure) DOCX exports. Bug fixed by using
                # the section-specific s3_key.
                artifact_key = block.get("s3_key", "")
                embedded = False
                if artifact_key:
                    try:
                        body = _s3().get_object(Bucket=S3_BUCKET_NAME, Key=artifact_key)["Body"].read()
                        lower = artifact_key.lower()
                        if lower.endswith((".png", ".jpg", ".jpeg")):
                            # python-docx natively embeds PNG/JPEG bytes.
                            doc.add_picture(io.BytesIO(body), width=Inches(6.5))
                            embedded = True
                        elif lower.endswith(".svg"):
                            if cairosvg:
                                png_buf = io.BytesIO()
                                cairosvg.svg2png(bytestring=body, write_to=png_buf, output_width=900)
                                png_buf.seek(0)
                                doc.add_picture(png_buf, width=Inches(6.5))
                                embedded = True
                            else:
                                logger.warning(
                                    f"[SAD] cairosvg unavailable; cannot embed SVG diagram {artifact_key}"
                                )
                        else:
                            logger.warning(
                                f"[SAD] unsupported diagram extension for DOCX: {artifact_key}"
                            )
                    except Exception as e:
                        logger.warning(
                            f"[SAD] failed to embed diagram {artifact_key} in DOCX: {e}"
                        )
                if not embedded:
                    doc.add_paragraph(
                        "[diagram unavailable â€” open the SAD in the app to "
                        "render it, then re-export the DOCX]"
                    )
        doc.add_paragraph("")  # spacer

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)

    headers = {
        "Content-Disposition": f'attachment; filename="SAD_{session_id}.docx"',
    }
    return StreamingResponse(
        out,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )
