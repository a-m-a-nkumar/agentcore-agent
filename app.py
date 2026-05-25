import os
import uuid
import hashlib
import json
import re
import asyncio
import time
import boto3
from botocore.exceptions import ClientError
from fastapi import FastAPI, UploadFile, File, HTTPException, Form, Depends, Header
from pydantic import BaseModel
from fastapi.responses import HTMLResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from fastapi.requests import Request
from fastapi.middleware.cors import CORSMiddleware
from dotenv import load_dotenv
from docx import Document
import io
import jwt
from jwt import PyJWKClient
from functools import wraps
from typing import Optional, List
import requests
from datetime import datetime

# Import API routers
from routers.projects import router as projects_router
from routers.sessions import router as sessions_router
from routers.integrations import router as integrations_router
from routers.integrations_internal import router as integrations_internal_router
from routers.sync import router as sync_router
from routers.jira_generation import router as jira_generation_router
from routers.orchestration import router as orchestration_router
from routers.orchestration_internal import router as orchestration_internal_router
from routers.test_generation import router as test_generation_router
from routers.brd_comparison import router as brd_comparison_router
from routers.test_internal import router as test_internal_router
from routers.design import router as design_router
from routers.design_sessions import router as design_sessions_router
from routers.sad import router as sad_router
from routers.brd import router as brd_router
from routers.harness import router as harness_router
from routers.pipeline_generator import router as pipeline_generator_router
from routers.terraform_generator import router as terraform_generator_router
from routers.figma import router as figma_router
# Import database helpers for session persistence
from db_helper import save_project_brd_session, create_or_update_user, update_user_access_role
# Environment-specific S3 implementation (local: plain boto3 | VDI: SSE-KMS)
from environment import s3_put_object, get_s3_client  # noqa: F401

load_dotenv(override=True)

# When using AWS_PROFILE (SSO), clear any stale STS credentials from the environment
# so boto3 resolves credentials via the SSO profile instead of expired keys
if os.getenv("AWS_PROFILE"):
    for key in ("AWS_ACCESS_KEY_ID", "AWS_SECRET_ACCESS_KEY", "AWS_SESSION_TOKEN"):
        os.environ.pop(key, None)

app = FastAPI(root_path=os.getenv("ROOT_PATH", ""))

# Load BMAD config (optional prompt overlay)
def _load_bmad_config():
    try:
        config_path = os.path.join(os.path.dirname(__file__), "bmad_agent_config.json")
        with open(config_path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return None

BMAD_CONFIG = _load_bmad_config()

def _build_bmad_prompt(base_prompt: str, workflow_key: str = "create-prd") -> str:
    """
    Prepend BMAD persona/principles/workflow prompt to the base prompt if config is available.
    Falls back to the base prompt unchanged if config is missing.
    """
    if not BMAD_CONFIG:
        return base_prompt

    persona = BMAD_CONFIG.get("agent", {}).get("persona", "").strip()
    principles = BMAD_CONFIG.get("agent", {}).get("principles", [])
    workflow = BMAD_CONFIG.get("workflows", {}).get(workflow_key, {})
    workflow_prompt = workflow.get("prompt", "").strip()

    parts = []
    if persona:
        parts.append(persona)
    if principles:
        parts.append("PRINCIPLES:\n" + "\n".join(f"- {p}" for p in principles))
    if workflow_prompt:
        parts.append(f"WORKFLOW: {workflow_key}\n{workflow_prompt}")
    parts.append(base_prompt)

    return "\n\n".join(parts)

# Add CORS middleware. The localhost origins cover the typical Vite + CRA dev
# servers. The two deployed frontend hosts are included so a developer can
# point the deployed UI at their LOCAL backend (running this file directly)
# during integration testing — without that, the browser blocks the cross-
# origin POST with "Response to preflight request doesn't pass access control
# check: No 'Access-Control-Allow-Origin' header is present".
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        # Local dev servers
        "http://localhost:8080", "http://localhost:8081", "http://localhost:5173",
        "http://127.0.0.1:8080", "http://127.0.0.1:8081", "http://127.0.0.1:5173",
        # Deployed frontends — needed when testing local backend against
        # the deployed UI in the browser.
        "https://sdlc-dev.deluxe.com",
        "https://ai-labs.deluxe.com",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*", "Authorization", "Content-Type"],
    expose_headers=["*"],
)

# Register API routers
app.include_router(projects_router)
app.include_router(sessions_router)
app.include_router(integrations_router)
app.include_router(integrations_internal_router)
app.include_router(sync_router)
app.include_router(orchestration_router)
app.include_router(orchestration_internal_router)
app.include_router(jira_generation_router)
app.include_router(test_generation_router)
app.include_router(brd_comparison_router)
app.include_router(test_internal_router)
app.include_router(design_router)
app.include_router(design_sessions_router)
app.include_router(sad_router)
# Unified BRD agent router (features/aman). Gated by BRD_USE_UNIFIED_AGENT
# in subsequent shim code -- registering the router itself is harmless
# because every endpoint requires auth + a session lookup that only
# succeeds for users opted into the unified path.
app.include_router(brd_router)
app.include_router(harness_router)
app.include_router(pipeline_generator_router)
app.include_router(terraform_generator_router)
app.include_router(figma_router)

# Add request logging middleware
@app.middleware("http")
async def log_requests(request: Request, call_next):
    """Log all incoming requests for debugging"""
    if (request.url.path.startswith("/api/upload-transcript") or
        request.url.path.startswith("/api/chat") or
        request.url.path.startswith("/api/analyst-chat") or
        request.url.path.startswith("/api/analyst-chat-stream")):
        auth_header = request.headers.get("authorization") or request.headers.get("Authorization")
        print(f"\n[REQUEST] {request.method} {request.url.path}")
        print(f"[REQUEST] Authorization header present: {bool(auth_header)}")
        if auth_header:
            print(f"[REQUEST] Auth header (first 30 chars): {auth_header[:30]}...")
        else:
            print(f"[REQUEST] All headers: {list(request.headers.keys())}")
    
    response = await call_next(request)
    return response

# Configuration (from .env / environment switch)
# ARN defaults come from environment.py — local account (448049797912) or VDI account (590184044598)
from environment import (  # noqa: E402
    DEFAULT_AGENT_ARN, DEFAULT_ANALYST_AGENT_ARN,
    DEFAULT_LAMBDA_BRD_CHAT, DEFAULT_LAMBDA_BRD_FROM_HISTORY,
    DEFAULT_LAMBDA_REQUIREMENTS_GATHERING, DEFAULT_LAMBDA_BRD_GENERATOR,
    DEFAULT_LAMBDA_REQUIREMENTS_GATHERING_ARN, DEFAULT_LAMBDA_BRD_FROM_HISTORY_ARN,
    DEFAULT_AGENTCORE_MEMORY_ID, DEFAULT_AGENTCORE_ACTOR_ID,
    S3_BUCKET_NAME,
)
AGENT_ARN = os.getenv("AGENT_ARN", DEFAULT_AGENT_ARN)
ANALYST_AGENT_ARN = os.getenv("ANALYST_AGENT_ARN", DEFAULT_ANALYST_AGENT_ARN)
REGION = os.getenv("AWS_REGION", "us-east-1")
LAMBDA_BRD_CHAT = os.getenv("LAMBDA_BRD_CHAT", DEFAULT_LAMBDA_BRD_CHAT)
LAMBDA_BRD_FROM_HISTORY = os.getenv("LAMBDA_BRD_FROM_HISTORY", DEFAULT_LAMBDA_BRD_FROM_HISTORY)
LAMBDA_REQUIREMENTS_GATHERING = os.getenv("LAMBDA_REQUIREMENTS_GATHERING", DEFAULT_LAMBDA_REQUIREMENTS_GATHERING)
LAMBDA_BRD_GENERATOR = os.getenv("LAMBDA_BRD_GENERATOR", DEFAULT_LAMBDA_BRD_GENERATOR)
AGENTCORE_MEMORY_ID = os.getenv("AGENTCORE_MEMORY_ID", DEFAULT_AGENTCORE_MEMORY_ID)
AGENTCORE_ACTOR_ID = os.getenv("AGENTCORE_ACTOR_ID", DEFAULT_AGENTCORE_ACTOR_ID)

# Log agent ARNs on startup
print(f"\n[CONFIG] Agent ARN: {AGENT_ARN}")
print(f"[CONFIG] Analyst Agent ARN: {ANALYST_AGENT_ARN}")
print(f"[CONFIG] Region: {REGION}\n")

# Azure AD Configuration
AZURE_CLIENT_ID = os.getenv("AZURE_CLIENT_ID", "")
AZURE_TENANT_ID = os.getenv("AZURE_TENANT_ID", "")
AZURE_CLIENT_SECRET = os.getenv("AZURE_CLIENT_SECRET", "")

# Import authentication functions
from auth import (
    verify_azure_token,
    store_user_identity_in_agentcore,
    get_user_identity_arn,
    check_brd_access_via_agentcore,
    grant_brd_access_via_agentcore,
    revoke_brd_access_via_agentcore,
    extract_user_groups,
    compute_allowed_modules,
    compute_access_role,
    require_module,
    GraphResolutionError,
)

# In-process cache: avoid hitting the DB on every authenticated request when
# the role hasn't changed. Per-worker dict; eventual consistency across
# workers is fine because the DB UPDATE itself is idempotent (WHERE
# access_role <> new_value short-circuits no-ops).
_LAST_ACCESS_ROLE_CACHE: dict[str, str] = {}

# Setup templates
templates = Jinja2Templates(directory="templates")

# Exception handler for validation errors
from fastapi.exceptions import RequestValidationError
from fastapi.responses import JSONResponse

@app.exception_handler(RequestValidationError)
async def validation_exception_handler(request: Request, exc: RequestValidationError):
    """Log detailed validation errors"""
    print(f"\n[VALIDATION ERROR] Path: {request.url.path}")
    print(f"[VALIDATION ERROR] Details: {json.dumps(exc.errors(), indent=2)}")
    try:
        body = await request.json()
        print(f"[VALIDATION ERROR] Body: {json.dumps(body, indent=2)}")
    except Exception:
        print(f"[VALIDATION ERROR] Body: <could not parse json>")
        
    return JSONResponse(
        status_code=422,
        content={"detail": exc.errors(), "body": "Check server logs for details"},
    )

# Helper function to check AWS credentials
def check_aws_credentials():
    """Check if AWS credentials are valid"""
    try:
        sts_client = boto3.client('sts', region_name=REGION)
        identity = sts_client.get_caller_identity()
        return True, identity
    except Exception as e:
        return False, str(e)

def get_agent_core_client():
    """Get a fresh AgentCore client with increased timeout for long-running operations"""
    from botocore.config import Config
    # Increase timeout to 5 minutes (300 seconds) for BRD generation
    config = Config(
        read_timeout=300,
        connect_timeout=10,
        retries={'max_attempts': 3}
    )
    return boto3.client('bedrock-agentcore', region_name=REGION, config=config)

def get_lambda_client():
    """Get a fresh Lambda client with extended timeout for long-running Lambda functions"""
    from botocore.config import Config
    # Increase timeout to 15 minutes (900 seconds) - max Lambda execution time
    config = Config(
        read_timeout=900,
        connect_timeout=60,
        retries={'max_attempts': 0}  # Don't retry on timeout - Lambda is already processing
    )
    return boto3.client('lambda', region_name=REGION, config=config)

def get_agentcore_identity_client():
    """Get AgentCore Identity client"""
    return boto3.client('bedrock-agentcore', region_name=REGION)

# -------------------------
# Azure AD Token Verification
# -------------------------



# -------------------------
# Authentication Decorator
# -------------------------

async def get_current_user(request: Request) -> dict:
    """FastAPI dependency to get current authenticated user"""
    # Get authorization header (case-insensitive)
    authorization = request.headers.get("authorization") or request.headers.get("Authorization")
    
    if not authorization:
        print(f"[AUTH] Authorization header missing. Headers: {list(request.headers.keys())}")
        raise HTTPException(status_code=401, detail="Authorization header missing")
    
    if not authorization.startswith("Bearer "):
        print(f"[AUTH] Invalid authorization header format: {authorization[:20]}...")
        raise HTTPException(status_code=401, detail="Invalid authorization header format")
    
    token = authorization.replace("Bearer ", "").strip()
    print(f"[AUTH] Token received (first 20 chars): {token[:20]}...")
    
    try:
        user_info = verify_azure_token(token)
        print(f"[AUTH] Token verified successfully for user: {user_info.get('email') or user_info.get('preferred_username')}")
    except HTTPException as e:
        print(f"[AUTH] Token verification failed: {e.detail}")
        raise
    except Exception as e:
        print(f"[AUTH] Unexpected error during token verification: {str(e)}")
        raise HTTPException(status_code=401, detail=f"Token verification failed: {str(e)}")
    
    user_id = user_info.get("oid") or user_info.get("sub")
    email = user_info.get("email") or user_info.get("preferred_username")
    name = user_info.get("name")

    print(f"[AUTH] User ID: {user_id}, Email: {email}")

    try:
        groups = extract_user_groups(user_info)
    except GraphResolutionError as e:
        # Overage user (>200 groups) whose Graph fallback failed. Do NOT
        # silently return empty modules — that would render AccessDenied
        # (a lie). 503 lets the frontend retry; transient Graph hiccups
        # then self-heal without anyone seeing the permanent denied page.
        print(f"[AUTH] Graph resolution failed for {email}: {e}")
        raise HTTPException(
            status_code=503,
            detail="Permission check temporarily unavailable — please retry in a moment.",
        )
    allowed_modules = compute_allowed_modules(groups)
    access_role = compute_access_role(groups)
    print(f"[AUTH] Groups: {groups}, Allowed modules: {allowed_modules}, access_role: {access_role}")

    # Persist access_role to users.access_role (cached per worker so we only
    # write when the role actually changes for this user).
    #
    # We pass email + name from the verified JWT so the UPSERT's INSERT
    # branch can satisfy users.email NOT NULL on a brand-new (or admin-
    # deleted) row. Without these, PostgreSQL's NOT NULL check on the
    # prospective INSERT row would fire BEFORE ON CONFLICT detection,
    # blocking the write even when the row already exists — that was the
    # "stuck NONE / NO GROUPS" pill scenario.
    if user_id and _LAST_ACCESS_ROLE_CACHE.get(user_id) != access_role:
        try:
            if update_user_access_role(user_id, access_role, email=email, name=name):
                _LAST_ACCESS_ROLE_CACHE[user_id] = access_role
        except Exception as e:
            print(f"[AUTH] Warning: Failed to persist access_role for {user_id}: {e}")

    # Store user identity in AgentCore if not exists
    try:
        store_user_identity_in_agentcore(user_id=user_id, email=email, name=name or "")
    except Exception as e:
        print(f"[AUTH] Warning: Failed to store user identity in AgentCore: {e}")

    return {
        "user_id": user_id,
        "email": email,
        "name": name,
        "token": token,
        "groups": groups,
        "allowed_modules": allowed_modules,
    }

def render_brd_json_to_text(brd_data: dict) -> str:
    """Render structured BRD JSON into readable plain text.
    Skips # In Scope and # Out of Scope as separate sections - they are subsections of Scope."""
    # Check if BRD uses sections format (newer format)
    if "sections" in brd_data:
        sections = brd_data.get("sections", [])
        lines = []
        lines.append("Business Requirements Document (BRD)")
        lines.append("")

        has_doc_title = False
        start_idx = 0
        if sections:
            first_title = (sections[0].get("title", "") or "").lower()
            if "ai-powered" in first_title or "brd" in first_title or (
                len(first_title) < 30 and not re.match(r'^\d+\.', first_title)
            ):
                has_doc_title = True
                start_idx = 1
                lines.append(sections[0].get("title", ""))
                lines.append("")

        section_counter = 1
        idx = start_idx
        while idx < len(sections):
            section = sections[idx]
            title = section.get("title", f"Section {section_counter}")
            title_lower = (title or "").lower().strip()
            if title.strip().startswith("#") and ("in scope" in title_lower or "out of scope" in title_lower):
                idx += 1
                continue
            title_clean = re.sub(r'^\d+\.\s*', '', title).strip()
            lines.append(f"{section_counter}. {title_clean}")
            lines.append("")

            content_blocks = list(section.get("content", []))
            if "scope" in title_lower and (not content_blocks or len(content_blocks) == 0):
                for i in (1, 2):
                    sub_idx = idx + i
                    if sub_idx < len(sections):
                        sub = sections[sub_idx]
                        sub_title = (sub.get("title", "") or "").lower()
                        if "# in scope" in sub_title:
                            content_blocks.append({"type": "paragraph", "text": "### In Scope"})
                            content_blocks.extend(sub.get("content", []))
                        elif "# out of scope" in sub_title:
                            content_blocks.append({"type": "paragraph", "text": "### Out of Scope"})
                            content_blocks.extend(sub.get("content", []))

            for block in content_blocks:
                block_type = block.get("type") if isinstance(block, dict) else None
                if block_type == "paragraph":
                    lines.append(block.get("text", "").strip())
                    lines.append("")
                elif block_type == "bullet":
                    for item in block.get("items", []):
                        lines.append(f"- {item}")
                    lines.append("")
                elif block_type == "table":
                    rows = block.get("rows", [])
                    if rows:
                        header = rows[0]
                        header_line = " | ".join(str(col) for col in header)
                        lines.append(header_line)
                        lines.append("-" * len(header_line))
                        for row in rows[1:]:
                            lines.append(" | ".join(str(col) for col in row))
                    lines.append("")
            section_counter += 1
            idx += 1
        return "\n".join(line.rstrip() for line in lines).rstrip() + "\n"
    
    # Fallback: Try to render as plain text if it's already text
    if isinstance(brd_data, str):
        return brd_data
    
    # Fallback: Convert to JSON string if all else fails
    return json.dumps(brd_data, indent=2, ensure_ascii=False)

def clean_markdown_text(text: str) -> str:
    """Remove markdown syntax from text"""
    if not text:
        return ""
    
    # Remove markdown headers (# ## ###) - but preserve the text after
    # Match: # Text or ## Text or ### Text
    text = re.sub(r'^#+\s+', '', text, flags=re.MULTILINE)
    
    # Remove horizontal rules (---) on their own line
    text = re.sub(r'^---+$', '', text, flags=re.MULTILINE)
    
    # Remove bold markdown (**text** or __text__) - handle nested cases
    # Match **text** but not ***text*** (that's bold+italic)
    text = re.sub(r'\*\*([^*]+?)\*\*', r'\1', text)
    text = re.sub(r'__([^_]+?)__', r'\1', text)
    
    # Remove italic markdown (*text* or _text_) - but preserve list markers (- item)
    # Only match if not at start of line with space after
    text = re.sub(r'(?<!^)(?<!\n)(?<!\s)\*([^*\n\s]+?)\*(?!\s)', r'\1', text)
    text = re.sub(r'(?<!^)(?<!\n)(?<!\s)_([^_\n\s]+?)_(?!\s)', r'\1', text)
    
    # Remove code blocks
    text = re.sub(r'```[\s\S]*?```', '', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    
    # Remove markdown table separators (|---|---| or |---|)
    text = re.sub(r'^\|?[\s\-|:]+\|?\s*$', '', text, flags=re.MULTILINE)
    
    # Clean up extra whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = text.strip()
    
    return text

def parse_markdown_table(text: str):
    """Parse markdown table format into rows"""
    import re
    lines = text.strip().split('\n')
    rows = []
    
    for line in lines:
        line = line.strip()
        if not line or line.startswith('|---'):
            continue
        
        # Split by | and clean up
        cells = [cell.strip() for cell in line.split('|')]
        # Remove empty first/last cells from split
        cells = [c for c in cells if c]
        if cells:
            rows.append(cells)
    
    return rows if rows else None

def render_brd_json_to_docx(brd_data) -> bytes:
    """Render structured BRD JSON or text into DOCX format with clean formatting
    
    Args:
        brd_data: Can be a dict (JSON structure) or str (plain text)
    """
    doc = Document()
    
    # Add title
    doc.add_heading('Business Requirements Document (BRD)', 0)
    
    # Check if BRD uses sections format (newer format)
    if isinstance(brd_data, dict) and "sections" in brd_data:
        sections = brd_data.get("sections", [])
        
        for idx, section in enumerate(sections, start=1):
            # Add section title as heading (clean markdown)
            section_title = section.get("title", f"Section {idx}")
            section_title = clean_markdown_text(section_title)
            doc.add_heading(section_title, level=1)
            
            # Process content blocks
            for block in section.get("content", []):
                block_type = block.get("type")
                
                if block_type == "paragraph":
                    # Add paragraph text (clean markdown headers but keep formatting)
                    text = block.get("text", "").strip()
                    if text:
                        # Clean headers only
                        text = re.sub(r'^#+\s+', '', text, flags=re.MULTILINE)
                        
                        # Check if paragraph contains markdown table (multi-line)
                        if '\n' in text and '|' in text and text.count('|') >= 2:
                            # Multi-line table in paragraph
                            lines_in_text = text.split('\n')
                            table_lines = []
                            regular_lines = []
                            
                            for txt_line in lines_in_text:
                                if '|' in txt_line and txt_line.count('|') >= 2 and not txt_line.strip().startswith('---'):
                                    table_lines.append(txt_line)
                                elif not txt_line.strip().startswith('---'):
                                    regular_lines.append(txt_line)
                            
                            # Process table if found
                            if table_lines:
                                table_data = parse_markdown_table('\n'.join(table_lines))
                                if table_data and len(table_data) > 0:
                                    max_cols = max(len(row) for row in table_data)
                                    table = doc.add_table(rows=len(table_data), cols=max_cols)
                                    table.style = 'Light Grid Accent 1'
                                    for row_idx, row_data in enumerate(table_data):
                                        for col_idx, cell_data in enumerate(row_data):
                                            if col_idx < len(table.rows[row_idx].cells):
                                                cell = table.rows[row_idx].cells[col_idx]
                                                cell.text = clean_markdown_text(str(cell_data))
                                    if len(table_data) > 0:
                                        header_cells = table.rows[0].cells
                                        for cell in header_cells:
                                            for paragraph in cell.paragraphs:
                                                for run in paragraph.runs:
                                                    run.bold = True
                            
                            # Process regular lines
                            for reg_line in regular_lines:
                                cleaned = clean_markdown_text(reg_line)
                                if cleaned:
                                    doc.add_paragraph(cleaned)
                        # Check if it's a single-line markdown table
                        elif '|' in text and text.count('|') >= 2:
                            table_rows = parse_markdown_table(text)
                            if table_rows and len(table_rows) > 0:
                                # Create Word table
                                max_cols = max(len(row) for row in table_rows)
                                table = doc.add_table(rows=len(table_rows), cols=max_cols)
                                table.style = 'Light Grid Accent 1'
                                
                                for row_idx, row_data in enumerate(table_rows):
                                    for col_idx, cell_data in enumerate(row_data):
                                        if col_idx < len(table.rows[row_idx].cells):
                                            cell = table.rows[row_idx].cells[col_idx]
                                            # Clean markdown from cell text
                                            cell.text = clean_markdown_text(str(cell_data))
                                
                                # Make header row bold
                                if len(table_rows) > 0:
                                    header_cells = table.rows[0].cells
                                    for cell in header_cells:
                                        for paragraph in cell.paragraphs:
                                            for run in paragraph.runs:
                                                run.bold = True
                            else:
                                # Not a table, just clean text
                                cleaned_text = clean_markdown_text(text)
                                if cleaned_text:
                                    doc.add_paragraph(cleaned_text)
                        else:
                            # Regular paragraph, clean markdown
                            cleaned_text = clean_markdown_text(text)
                            if cleaned_text:
                                doc.add_paragraph(cleaned_text)
                
                elif block_type == "bullet":
                    # Add bullet list (preserve markdown formatting)
                    items = block.get("items", [])
                    if items:
                        for item in items:
                            cleaned_item = clean_markdown_text(str(item))
                            if cleaned_item:
                                doc.add_paragraph(cleaned_item, style='List Bullet')
                
                elif block_type == "table":
                    # Add table
                    rows = block.get("rows", [])
                    if rows:
                        # Create table with appropriate dimensions
                        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                        table.style = 'Light Grid Accent 1'
                        
                        # Populate table
                        for row_idx, row_data in enumerate(rows):
                            for col_idx, cell_data in enumerate(row_data):
                                if col_idx < len(table.rows[row_idx].cells):
                                    cell = table.rows[row_idx].cells[col_idx]
                                    # Clean markdown from cell text
                                    cell.text = clean_markdown_text(str(cell_data))
                        
                        # Make header row bold if it's the first row
                        if len(rows) > 0:
                            header_cells = table.rows[0].cells
                            for cell in header_cells:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
            
            # Add spacing between sections
            doc.add_paragraph("")
    
    # Fallback: If it's plain text, parse and clean it
    elif isinstance(brd_data, str):
        # Try to parse as markdown and convert
        lines = brd_data.split('\n')
        current_paragraph = []
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                if current_paragraph:
                    text = ' '.join(current_paragraph)
                    cleaned = clean_markdown_text(text)
                    if cleaned:
                        doc.add_paragraph(cleaned)
                    current_paragraph = []
                i += 1
                continue
            
            # Check for markdown table
            if '|' in line and line.count('|') >= 2 and not line.startswith('|---'):
                if current_paragraph:
                    text = ' '.join(current_paragraph)
                    cleaned = clean_markdown_text(text)
                    if cleaned:
                        doc.add_paragraph(cleaned)
                    current_paragraph = []
                
                # Collect table rows
                table_rows = []
                j = i
                while j < len(lines) and ('|' in lines[j] or lines[j].strip().startswith('---')):
                    if not lines[j].strip().startswith('---'):
                        table_rows.append(lines[j])
                    j += 1
                
                if table_rows:
                    table_data = parse_markdown_table('\n'.join(table_rows))
                    if table_data and len(table_data) > 0:
                        # Determine max columns
                        max_cols = max(len(row) for row in table_data)
                        table = doc.add_table(rows=len(table_data), cols=max_cols)
                        table.style = 'Light Grid Accent 1'
                        for row_idx, row_data in enumerate(table_data):
                            for col_idx, cell_data in enumerate(row_data):
                                if col_idx < len(table.rows[row_idx].cells):
                                    cell = table.rows[row_idx].cells[col_idx]
                                    # Clean markdown from cell text
                                    cell.text = clean_markdown_text(str(cell_data))
                        # Make header row bold
                        if len(table_data) > 0:
                            header_cells = table.rows[0].cells
                            for cell in header_cells:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                
                i = j  # Skip processed table lines
            else:
                # Regular line - check for markdown headers
                if line.startswith('#'):
                    # It's a header - add previous paragraph if any
                    if current_paragraph:
                        text = ' '.join(current_paragraph)
                        cleaned = clean_markdown_text(text)
                        if cleaned:
                            doc.add_paragraph(cleaned)
                        current_paragraph = []
                    
                    # Determine header level
                    header_level = 0
                    while header_level < len(line) and line[header_level] == '#':
                        header_level += 1
                    
                    # Extract header text
                    header_text = line[header_level:].strip()
                    cleaned_header = clean_markdown_text(header_text)
                    if cleaned_header:
                        # Use appropriate heading level (max level 3 for Word)
                        level = min(header_level, 3)
                        doc.add_heading(cleaned_header, level=level)
                else:
                    # Regular line
                    if line and not line.startswith('---'):
                        current_paragraph.append(line)
                i += 1
        
        # Add remaining paragraph
        if current_paragraph:
            text = ' '.join(current_paragraph)
            cleaned = clean_markdown_text(text)
            if cleaned:
                doc.add_paragraph(cleaned)
    
    # Save to bytes
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    return docx_bytes.read()

# Initialize clients on startup
try:
    # Verify credentials on startup
    creds_valid, creds_info = check_aws_credentials()
    if creds_valid:
        print(f"[APP] ✅ AWS credentials valid. Account: {creds_info.get('Account', 'Unknown')}")
        print(f"[APP] User: {creds_info.get('Arn', 'Unknown')}")
    else:
        print(f"[APP] [WARN] AWS credentials check failed: {creds_info}")
        print("[APP] Please configure AWS credentials using:")
        print("  - AWS CLI: aws configure")
        print("  - Environment variables: AWS_ACCESS_KEY_ID, AWS_SECRET_ACCESS_KEY")
        print("  - Or use AWS SSO/credentials file")
    
    # Test client initialization
    test_client = get_agent_core_client()
    print(f"[APP] ✅ AgentCore client initialized successfully")
except Exception as e:
    print(f"[APP] [ERROR] Failed to initialize AWS clients: {e}")
    print("[APP] Please configure AWS credentials using:")
    print("  - AWS CLI: aws configure")
    print("  - Environment variables: AWS_ACCESS_KEY_ID, AWS_SECRET_ACCESS_KEY")
    print("  - Or use AWS SSO/credentials file")

def read_docx(file_content):
    doc = Document(io.BytesIO(file_content))
    return "\n".join([p.text for p in doc.paragraphs])

def read_pdf(file_content):
    # Use `pypdf` (the modern fork) — it's what requirements.txt installs.
    # Importing `PyPDF2` instead would ImportError at runtime since that
    # package is not in the deployed image.
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(file_content))
    return "\n".join([page.extract_text() or "" for page in reader.pages])

def extract_text(file_content, filename):
    """Extract text from .docx, .pdf, or .txt files."""
    lower = filename.lower()
    if lower.endswith(".docx"):
        return read_docx(file_content)
    elif lower.endswith(".pdf"):
        return read_pdf(file_content)
    else:
        return file_content.decode("utf-8", errors="replace")

@app.get("/", response_class=HTMLResponse)
async def read_root(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})

@app.post("/generate")
async def generate_brd(
    transcript: UploadFile = File(...),
    template: UploadFile = File(...),
    current_user: dict = Depends(get_current_user)
):
    try:
        print("\n" + "="*80)
        print("[APP] Starting BRD generation")
        print("="*80)
        
        # 1. Read files
        transcript_content = await transcript.read()
        template_content = await template.read()
        
        print(f"[APP] Transcript file: {transcript.filename} ({len(transcript_content)} bytes)")
        print(f"[APP] Template file: {template.filename} ({len(template_content)} bytes)")
        
        # 2. Extract text (supports .docx, .pdf, .txt)
        transcript_text = extract_text(transcript_content, transcript.filename)

        template_text = read_docx(template_content)

        print(f"[APP] Transcript text: {len(transcript_text)} chars")
        print(f"[APP] Template text: {len(template_text)} chars")

        # Validate transcript is not empty
        if not transcript_text or len(transcript_text.strip()) < 50:
            return JSONResponse(status_code=400, content={
                "error": "Transcript is empty or too short. Please upload a transcript with meaningful content."
            })

        # 3. Prepare Payload (with BMAD persona/workflow overlay if available)
        base_prompt = "Generate a BRD based on the provided template and transcript."
        bmad_prompt = _build_bmad_prompt(base_prompt, workflow_key="create-prd")
        payload_dict = {
            "prompt": bmad_prompt,
            "template": template_text,
            "transcript": transcript_text,
            "user_id": current_user.get("user_id"),  # for token usage tracking
        }
        payload_bytes = json.dumps(payload_dict).encode('utf-8')
        
        print(f"[APP] Payload size: {len(payload_bytes)} bytes")
        
        # 4. Invoke Agent
        session_id = str(uuid.uuid4())
        print(f"[APP] Session ID: {session_id}")
        print(f"[APP] Agent ARN: {AGENT_ARN}")
        print(f"[APP] Calling agent...")
        print(f"[APP] Note: BRD generation may take 1-3 minutes. Please wait...")
        
        # Get fresh client to ensure we use latest credentials
        # Increased timeout to 5 minutes for BRD generation
        agent_core_client = get_agent_core_client()
        
        try:
            loop = asyncio.get_event_loop()
            response = await loop.run_in_executor(
                None,
                lambda: agent_core_client.invoke_agent_runtime(
                    agentRuntimeArn=AGENT_ARN,
                    runtimeSessionId=session_id,
                    payload=payload_bytes,
                    qualifier="DEFAULT"
                )
            )
        except Exception as timeout_error:
            if "timeout" in str(timeout_error).lower() or "ReadTimeoutError" in str(type(timeout_error).__name__):
                print(f"[APP] ⚠️  Request timed out. The agent may still be processing.")
                print(f"[APP] This can happen if the BRD is very large or the agent is slow.")
                print(f"[APP] Try checking CloudWatch logs or reducing the transcript/template size.")
                return JSONResponse(status_code=504, content={
                    "error": "Request timeout - agent took too long to respond",
                    "message": "BRD generation is taking longer than expected. The agent may still be processing. Try:\n1. Checking CloudWatch logs\n2. Reducing transcript/template size\n3. Retrying the request",
                    "type": "TimeoutError"
                })
            raise
        
        print(f"[APP] Agent response received")
        
        # 5. Parse Response
        content = []
        for chunk in response.get("response", []):
            content.append(chunk.decode('utf-8'))
            
        full_response_str = ''.join(content)
        
        print(f"[APP] Response length: {len(full_response_str)} chars")
        print(f"[APP] Response preview: {full_response_str[:300]}")
        
        # The agent now returns clean JSON with the BRD
        try:
            # First parse the outer response
            result_json = json.loads(full_response_str)
            print(f"[APP] Parsed as JSON, keys: {list(result_json.keys())}")
            
            # The result field contains the agent's response
            if 'result' in result_json:
                result_str = result_json['result']
                print(f"[APP] Result preview: {result_str[:200]}")
                
                # Agent now returns JSON with {status, brd, brd_id}
                try:
                    agent_data = json.loads(result_str)
                    print(f"[APP] Agent data keys: {list(agent_data.keys())}")
                    
                    if agent_data.get('brd'):
                        print(f"[APP] Found BRD! Length: {len(agent_data['brd'])} chars")
                        brd_id = agent_data.get('brd_id')
                        
                        # Create AgentCore Memory session for this BRD
                        session_id = None
                        if brd_id:
                            try:
                                print(f"[APP] Creating AgentCore Memory session for BRD {brd_id}")
                                # Call Lambda to create session
                                lambda_client = get_lambda_client()
                                session_payload = {
                                    'action': 'create_session',
                                    'brd_id': brd_id,
                                    'template': template_text[:500],  # Truncate for session creation
                                    'transcript': transcript_text[:500]  # Truncate for session creation
                                }
                                session_response = lambda_client.invoke(
                                    FunctionName=LAMBDA_BRD_CHAT,
                                    InvocationType='RequestResponse',
                                    Payload=json.dumps(session_payload)
                                )
                                session_result = json.loads(session_response['Payload'].read())
                                if session_result.get('statusCode') == 200:
                                    session_body = json.loads(session_result.get('body', '{}'))
                                    session_id = session_body.get('session_id')
                                    print(f"[APP] ✅ Created session: {session_id}")
                                else:
                                    print(f"[APP] ⚠️  Session creation failed, will auto-create on first chat")
                            except Exception as e:
                                print(f"[APP] ⚠️  Failed to create session: {e}, will auto-create on first chat")
                        
                        return JSONResponse(content={
                            'result': agent_data['brd'],
                            'brd_id': brd_id,
                            'session_id': session_id,  # Return session_id to frontend
                            'status': 'success'
                        })
                except json.JSONDecodeError:
                    # If result is not JSON, return as-is
                    pass
            
            return JSONResponse(content=result_json)
            
        except json.JSONDecodeError as e:
            print(f"[APP] JSON decode error: {e}")
            return JSONResponse(content={"result": full_response_str})

    except Exception as e:
        error_msg = str(e)
        print(f"[APP] ERROR: {error_msg}")
        import traceback
        traceback.print_exc()
        
        # Check if it's a credentials issue
        if "AccessDeniedException" in error_msg or "security token" in error_msg.lower() or "invalid" in error_msg.lower():
            creds_valid, creds_info = check_aws_credentials()
            if not creds_valid:
                error_msg = f"AWS credentials are invalid or expired. Please refresh your credentials.\n\nTo fix:\n1. Run: aws configure\n2. Or set environment variables: AWS_ACCESS_KEY_ID, AWS_SECRET_ACCESS_KEY\n3. Or refresh AWS SSO: aws sso login\n\nError details: {creds_info}"
            else:
                error_msg = f"AWS credentials are valid but access denied. Check IAM permissions for AgentCore.\n\nOriginal error: {error_msg}"
        
        return JSONResponse(status_code=500, content={
            "error": error_msg,
            "message": error_msg,
            "type": "AccessDeniedException" if "AccessDeniedException" in str(e) else "UnknownError"
        })

@app.post("/api/upload-transcript")
async def upload_transcript_to_s3(
    request: Request,
    transcripts: List[UploadFile] = File(...),
    current_user: dict = Depends(get_current_user)
):
    """Upload one or more transcript files to S3 and return S3 paths"""
    try:
        print("\n" + "="*80)
        print(f"[UPLOAD] Uploading {len(transcripts)} transcript(s) to S3")
        print(f"[UPLOAD] User: {current_user.get('email')} ({current_user.get('user_id')})")
        print("="*80)

        bucket_name = S3_BUCKET_NAME
        uploaded_files = []

        for transcript in transcripts:
            transcript_id = str(uuid.uuid4())
            transcript_key = f"transcripts/{transcript_id}/{transcript.filename}"
            transcript_content = await transcript.read()

            print(f"[UPLOAD] Uploading to S3: s3://{bucket_name}/{transcript_key}")
            print(f"[UPLOAD] File: {transcript.filename}, Size: {len(transcript_content)} bytes")

            s3_put_object(
                key=transcript_key,
                body=transcript_content,
                content_type=transcript.content_type or "application/octet-stream",
                bucket=bucket_name,
            )
            uploaded_files.append({
                "transcript_id": transcript_id,
                "s3_path": transcript_key,
                "s3_url": f"s3://{bucket_name}/{transcript_key}",
                "filename": transcript.filename
            })

        print(f"[UPLOAD] ✅ Successfully uploaded {len(uploaded_files)} file(s) to S3")

        # Return both multi-file format and single-file backward compat
        result = {
            "success": True,
            "files": uploaded_files,
        }
        # Backward compat: also set top-level fields from first file
        if uploaded_files:
            result["transcript_id"] = uploaded_files[0]["transcript_id"]
            result["s3_path"] = uploaded_files[0]["s3_path"]
            result["s3_url"] = uploaded_files[0]["s3_url"]
            result["filename"] = uploaded_files[0]["filename"]

        return JSONResponse(content=result)

    except Exception as e:
        error_msg = str(e)
        print(f"[UPLOAD] ERROR: {error_msg}")
        import traceback
        traceback.print_exc()
        return JSONResponse(status_code=500, content={
            "error": error_msg,
            "message": f"Failed to upload transcript to S3: {error_msg}"
        })

@app.post("/api/generate-from-s3")
async def generate_brd_from_s3(
    transcript_s3_paths: Optional[str] = Form(None),
    transcript_s3_path: Optional[str] = Form(None),
    project_id: Optional[str] = Form(None),
    current_user: dict = Depends(get_current_user)
):
    """Generate BRD from transcript(s) in S3 and template in S3"""
    t0 = time.time()
    try:
        print("\n" + "="*80)
        print("[APP] Starting BRD generation from S3")
        print("="*80)

        # Support both plural (new) and singular (backward compat) field names.
        # Filenames can contain commas (e.g. "Monday, March 30, part 2.txt"), so a
        # JSON array is the canonical form. Fall back to comma-split only when the
        # payload isn't valid JSON (older clients).
        raw_paths = (transcript_s3_paths or transcript_s3_path or "").strip()
        s3_paths: list[str] = []
        if raw_paths:
            try:
                parsed = json.loads(raw_paths)
                if isinstance(parsed, list):
                    s3_paths = [str(p).strip() for p in parsed if str(p).strip()]
                elif isinstance(parsed, str):
                    s3_paths = [parsed.strip()] if parsed.strip() else []
            except (ValueError, TypeError):
                s3_paths = [p.strip() for p in raw_paths.split(",") if p.strip()]
        if not s3_paths:
            return JSONResponse(status_code=400, content={"error": "No transcript S3 paths provided"})

        s3_client = get_s3_client()
        bucket_name = S3_BUCKET_NAME

        # Template path in S3
        template_s3_path_key = "templates/Deluxe_BRD_Template.docx"

        print(f"[APP] Transcript S3 paths ({len(s3_paths)}): {s3_paths}")
        print(f"[APP] Template S3 path: {template_s3_path_key}")

        # 1. Fetch and extract text from all transcript files
        transcript_texts = []
        for s3_path in s3_paths:
            print(f"[APP] Fetching transcript from S3: {s3_path}")
            resp = s3_client.get_object(Bucket=bucket_name, Key=s3_path)
            content = resp['Body'].read()
            print(f"[APP] File: {s3_path}, Size: {len(content)} bytes")
            text = extract_text(content, s3_path)
            transcript_texts.append(text)

        transcript_text = "\n\n---\n\n".join(transcript_texts)

        # 2. Fetch template from S3
        print(f"[APP] Fetching template from S3...")
        template_response = s3_client.get_object(Bucket=bucket_name, Key=template_s3_path_key)
        template_content = template_response['Body'].read()

        print(f"[APP] Combined transcript text: {len(transcript_text)} chars")
        print(f"[APP] Template file: {len(template_content)} bytes")

        # 3. Extract template text
        template_text = read_docx(template_content)

        print(f"[APP] Transcript text: {len(transcript_text)} chars")
        print(f"[APP] Template text: {len(template_text)} chars")

        # Validate transcript is not empty
        if not transcript_text or len(transcript_text.strip()) < 50:
            return JSONResponse(status_code=400, content={
                "error": "Transcript is empty or too short. Please upload a transcript with meaningful content."
            })

        # 4. Prepare Payload (same as /generate endpoint, with BMAD overlay if available)
        base_prompt = "Generate a BRD based on the provided template and transcript."
        bmad_prompt = _build_bmad_prompt(base_prompt, workflow_key="create-prd")
        payload_dict = {
            "prompt": bmad_prompt,
            "template": template_text,
            "transcript": transcript_text,
            "user_id": current_user.get("user_id"),  # for token usage tracking
        }
        payload_bytes = json.dumps(payload_dict).encode('utf-8')
        
        print(f"[APP] Payload size: {len(payload_bytes)} bytes")
        
        # 5. Invoke Agent (same as /generate endpoint)
        session_id = str(uuid.uuid4())
        print(f"[APP] Session ID: {session_id}")
        print(f"[APP] Agent ARN: {AGENT_ARN}")
        print(f"[APP] Calling agent...")
        print(f"[APP] Note: BRD generation may take 1-3 minutes. Please wait...")
        
        agent_core_client = get_agent_core_client()
        
        try:
            loop = asyncio.get_event_loop()
            response = await loop.run_in_executor(
                None,
                lambda: agent_core_client.invoke_agent_runtime(
                    agentRuntimeArn=AGENT_ARN,
                    runtimeSessionId=session_id,
                    payload=payload_bytes,
                    qualifier="DEFAULT"
                )
            )
        except Exception as timeout_error:
            if "timeout" in str(timeout_error).lower() or "ReadTimeoutError" in str(type(timeout_error).__name__):
                print(f"[APP] ⚠️  Request timed out. The agent may still be processing.")
                return JSONResponse(status_code=504, content={
                    "error": "Request timeout - agent took too long to respond",
                    "message": "BRD generation is taking longer than expected. The agent may still be processing.",
                    "type": "TimeoutError"
                })
            raise
        
        print(f"[APP] Agent response received")
        
        # 6. Parse Response (same as /generate endpoint)
        content = []
        for chunk in response.get("response", []):
            content.append(chunk.decode('utf-8'))
            
        full_response_str = ''.join(content)
        
        print(f"[APP] Response length: {len(full_response_str)} chars")
        print(f"[APP] Response preview: {full_response_str[:300]}")
        
        try:
            result_json = json.loads(full_response_str)
            print(f"[APP] Parsed as JSON, keys: {list(result_json.keys())}")
            
            if 'result' in result_json:
                result_str = result_json['result']
                print(f"[APP] Result preview: {result_str[:200]}")
                
                try:
                    agent_data = json.loads(result_str)
                    print(f"[APP] Agent data keys: {list(agent_data.keys())}")
                    
                    if agent_data.get('brd'):
                        print(f"[APP] Found BRD! Length: {len(agent_data['brd'])} chars")
                        brd_id = agent_data.get('brd_id')
                        
                        # Create AgentCore Memory session for this BRD
                        session_id_memory = None
                        if brd_id:
                            try:
                                print(f"[APP] Creating AgentCore Memory session for BRD {brd_id}")
                                lambda_client = get_lambda_client()
                                session_payload = {
                                    'action': 'create_session',
                                    'brd_id': brd_id,
                                    'template': template_text[:500],
                                    'transcript': transcript_text[:500]
                                }
                                session_response = lambda_client.invoke(
                                    FunctionName=LAMBDA_BRD_CHAT,
                                    InvocationType='RequestResponse',
                                    Payload=json.dumps(session_payload)
                                )
                                session_result = json.loads(session_response['Payload'].read())
                                if session_result.get('statusCode') == 200:
                                    session_body = json.loads(session_result.get('body', '{}'))
                                    session_id_memory = session_body.get('session_id')
                                    print(f"[APP] ✅ Created session: {session_id_memory}")
                                else:
                                    print(f"[APP] ⚠️  Session creation failed, will auto-create on first chat")
                            except Exception as e:
                                print(f"[APP] ⚠️  Failed to create session: {e}, will auto-create on first chat")
                        
                        # Persist BRD session to project for future restoration
                        if project_id and brd_id:
                            try:
                                save_project_brd_session(
                                    project_id=project_id,
                                    brd_id=brd_id,
                                    agentcore_session_id=session_id_memory
                                )
                                print(f"[APP] ✅ Persisted BRD session to project {project_id}")
                            except Exception as e:
                                print(f"[APP] ⚠️  Failed to persist BRD session: {e}")
                        
                        try:
                            from db_helper import track_event
                            track_event(
                                current_user["user_id"],
                                module="brd",
                                event_type="pm_agent_brd_generated",
                                project_id=project_id,
                                metadata={
                                    "transcript_s3_paths": s3_paths,
                                    "transcript_count": len(s3_paths),
                                    "brd_id": brd_id,
                                    "duration_ms": int((time.time() - t0) * 1000),
                                },
                            )
                        except Exception as _track_err:
                            print(f"[APP] track_event failed (non-fatal): {_track_err}")

                        return JSONResponse(content={
                            'result': agent_data['brd'],
                            'brd_id': brd_id,
                            'session_id': session_id_memory,
                            'status': 'success'
                        })
                except json.JSONDecodeError:
                    pass
            
            return JSONResponse(content=result_json)
            
        except json.JSONDecodeError as e:
            print(f"[APP] JSON decode error: {e}")
            return JSONResponse(content={"result": full_response_str})

    except Exception as e:
        error_msg = str(e)
        print(f"[APP] ERROR: {error_msg}")
        import traceback
        traceback.print_exc()
        
        if "AccessDeniedException" in error_msg or "security token" in error_msg.lower() or "invalid" in error_msg.lower():
            creds_valid, creds_info = check_aws_credentials()
            if not creds_valid:
                error_msg = f"AWS credentials are invalid or expired. Please refresh your credentials.\n\nTo fix:\n1. Run: aws configure\n2. Or set environment variables: AWS_ACCESS_KEY_ID, AWS_SECRET_ACCESS_KEY\n3. Or refresh AWS SSO: aws sso login\n\nError details: {creds_info}"
            else:
                error_msg = f"AWS credentials are valid but access denied. Check IAM permissions for AgentCore.\n\nOriginal error: {error_msg}"
        
        return JSONResponse(status_code=500, content={
            "error": error_msg,
            "message": error_msg,
            "type": "AccessDeniedException" if "AccessDeniedException" in str(e) else "UnknownError"
        })

@app.post("/api/chat")
async def chat_with_agent(
    message: str = Form(...),
    brd_id: str = Form(...),
    session_id: str = Form(...),
    current_user: dict = Depends(get_current_user)
):
    try:
        print(f"\n[CHAT] Message: {message}")
        print(f"[CHAT] BRD ID: {brd_id}")
        print(f"[CHAT] Session ID: {session_id} (length: {len(session_id)})")
        
        # Ensure brd_id is valid (not "none")
        if brd_id == "none" or not brd_id:
            return JSONResponse(status_code=400, content={
                "error": "BRD ID is required for chat. Please generate a BRD first.",
                "result": "Error: No BRD ID provided. Please upload a transcript and generate a BRD first."
            })
        
        # Ensure session_id is valid (not "none")
        if session_id == "none" or not session_id:
            # Generate a session ID based on BRD ID for consistency
            session_id = f"brd-session-{brd_id}"
            print(f"[CHAT] Session ID was 'none', generated: {session_id}")
        
        # Format the message to be clear for the agent
        # Include session_id in the payload so the agent can pass it to the Lambda
        # The agent entrypoint will extract brd_id and pass it to chat_with_brd tool
        formatted_message = message.strip()
        
        # Include session_id in payload so agent can use it when calling chat_with_brd
        payload_dict = {
            "prompt": formatted_message,
            "brd_id": brd_id,
            "session_id": session_id,  # Pass session_id so agent can use it
            "user_id": current_user.get("user_id"),  # for token usage tracking
        }
        payload_bytes = json.dumps(payload_dict).encode('utf-8')
        
        print(f"[CHAT] Payload: {payload_dict}")
        print(f"[CHAT] Calling agent...")
        
        # Get fresh client to ensure we use latest credentials
        agent_core_client = get_agent_core_client()
        
        # Use the consistent session_id as runtimeSessionId so that conversation
        # events are stored in AgentCore Memory under this session and can be
        # retrieved later by /api/brd-history.
        print(f"[CHAT] Using runtimeSessionId: {session_id}")

        loop = asyncio.get_event_loop()
        response = await loop.run_in_executor(
            None,
            lambda: agent_core_client.invoke_agent_runtime(
                agentRuntimeArn=AGENT_ARN,
                runtimeSessionId=session_id,
                payload=payload_bytes,
                qualifier="DEFAULT"
            )
        )

        content = []
        for chunk in response.get("response", []):
            content.append(chunk.decode('utf-8'))
            
        full_response_str = ''.join(content)
        print(f"[CHAT] Raw response: {full_response_str[:500]}")
        print(f"[CHAT] Raw response type: {type(full_response_str)}")
        print(f"[CHAT] Raw response length: {len(full_response_str)}")
        
        # Parse the agent response to extract the actual text content
        final_text = None
        try:
            result_json = json.loads(full_response_str)
            print(f"[CHAT] Parsed JSON, keys: {list(result_json.keys()) if isinstance(result_json, dict) else 'Not a dict'}")
            print(f"[CHAT] Parsed JSON type: {type(result_json)}")

            # AgentCore returns responses in format: {'role': 'assistant', 'content': [{'text': '...'}]}
            # Extract the text from the content array
            extracted_text = None

            if isinstance(result_json, dict):
                # Check for 'content' field with text
                if 'content' in result_json:
                    content_list = result_json['content']
                    if isinstance(content_list, list) and len(content_list) > 0:
                        first_content = content_list[0]
                        if isinstance(first_content, dict) and 'text' in first_content:
                            extracted_text = first_content['text']

                # Also check for 'result' field (some responses use this)
                if not extracted_text and 'result' in result_json:
                    result_value = result_json['result']
                    if isinstance(result_value, str):
                        extracted_text = result_value
                    elif isinstance(result_value, dict):
                        # Try to extract from nested result
                        if 'content' in result_value:
                            content_list = result_value['content']
                            if isinstance(content_list, list) and len(content_list) > 0:
                                first_content = content_list[0]
                                if isinstance(first_content, dict) and 'text' in first_content:
                                    extracted_text = first_content['text']

                # Check for direct 'text' or 'message' fields
                if not extracted_text:
                    extracted_text = result_json.get('text') or result_json.get('message') or result_json.get('response')

            # Determine final_text from extraction results
            if extracted_text:
                print(f"[CHAT] ✅ Extracted text successfully: {extracted_text[:200]}")
                print(f"[CHAT] Extracted text type: {type(extracted_text)}")
                print(f"[CHAT] Extracted text length: {len(extracted_text)}")
                if not isinstance(extracted_text, str):
                    extracted_text = str(extracted_text)
                final_text = extracted_text
            else:
                print(f"[CHAT] Could not extract text, trying to format response")
                if isinstance(result_json, dict):
                    for key in ['text', 'message', 'content', 'result', 'response', 'answer']:
                        if key in result_json:
                            value = result_json[key]
                            if isinstance(value, str) and value.strip():
                                final_text = value
                                break
                            elif isinstance(value, list) and len(value) > 0:
                                if isinstance(value[0], dict) and 'text' in value[0]:
                                    final_text = value[0]['text']
                                    break

                if not final_text:
                    final_text = json.dumps(result_json, indent=2) if isinstance(result_json, dict) else full_response_str

        except json.JSONDecodeError:
            print(f"[CHAT] Response is not JSON, returning as text")
            final_text = full_response_str

        return JSONResponse(content={
            "result": final_text,
            "response": final_text,
            "session_id": session_id
        })

    except Exception as e:
        error_msg = str(e)
        print(f"[CHAT] ERROR: {error_msg}")
        import traceback
        traceback.print_exc()
        
        # Check if it's a credentials issue
        if "AccessDeniedException" in error_msg or "security token" in error_msg.lower() or "invalid" in error_msg.lower():
            creds_valid, creds_info = check_aws_credentials()
            if not creds_valid:
                error_msg = f"AWS credentials are invalid or expired. Please refresh your credentials.\n\nTo fix:\n1. Run: aws configure\n2. Or set environment variables: AWS_ACCESS_KEY_ID, AWS_SECRET_ACCESS_KEY\n3. Or refresh AWS SSO: aws sso login\n\nError details: {creds_info}"
            else:
                error_msg = f"AWS credentials are valid but access denied. Check IAM permissions for AgentCore.\n\nOriginal error: {error_msg}"
        
        return JSONResponse(status_code=500, content={
            "error": error_msg,
            "result": f"Error: {error_msg}",
            "type": "AccessDeniedException" if "AccessDeniedException" in str(e) else "UnknownError"
        })

def extract_text_from_analyst_response(response_str: str) -> tuple[str, str]:
    """
    Extract plain text message and session_id from analyst agent's JSON response.
    Handles both direct analyst agent responses and AgentCore-wrapped responses.
    Returns: (message_text, session_id)
    """
    print(f"[extract_text] Called with response length: {len(response_str) if response_str else 0}")
    
    if not response_str or not isinstance(response_str, str):
        print(f"[extract_text] Response is None or not string")
        return None, None
    
    response_trimmed = response_str.strip()
    if not response_trimmed.startswith('{'):
        print(f"[extract_text] Response doesn't start with '{{', returning as plain text")
        return response_str, None
    
    print(f"[extract_text] Response starts with '{{'', attempting JSON parse")
    print(f"[extract_text] First 200 chars: {response_trimmed[:200]}")
    
    try:
        parsed = json.loads(response_trimmed)
        print(f"[extract_text] JSON parse successful, type: {type(parsed)}")
        
        if isinstance(parsed, dict):
            print(f"[extract_text] Parsed dict keys: {list(parsed.keys())}")
            
            # Case 1: Direct analyst agent response: {"result": "...", "session_id": "...", "message": "..."}
            has_message = 'message' in parsed
            has_result_and_session = ('result' in parsed and 'session_id' in parsed)
            print(f"[extract_text] has_message: {has_message}, has_result_and_session: {has_result_and_session}")
            
            if has_message or has_result_and_session:
                message_text = parsed.get('message') or parsed.get('result')
                session_id = parsed.get('session_id')
                print(f"[extract_text] Extracted message_text type: {type(message_text)}, length: {len(message_text) if isinstance(message_text, str) else 'N/A'}")
                print(f"[extract_text] Extracted session_id: {session_id}")
                
                if message_text and isinstance(message_text, str):
                    print(f"[extract_text] ✅ Returning message_text (string)")
                    return message_text, session_id
                elif message_text:
                    # If message_text is not a string, convert it
                    print(f"[extract_text] ✅ Returning message_text (converted to string)")
                    return str(message_text), session_id
                else:
                    print(f"[extract_text] ⚠️ message_text is None or empty")
            
            # Case 2: AgentCore wrapped response: {"result": "{\"result\": \"...\", \"session_id\": \"...\", \"message\": \"...\"}"}
            if 'result' in parsed:
                result_value = parsed.get('result')
                if isinstance(result_value, str) and result_value.strip().startswith('{'):
                    # Try to parse the nested JSON string
                    try:
                        nested_parsed = json.loads(result_value)
                        if isinstance(nested_parsed, dict):
                            nested_message = nested_parsed.get('message') or nested_parsed.get('result')
                            nested_session_id = nested_parsed.get('session_id')
                            if nested_message and isinstance(nested_message, str):
                                return nested_message, nested_session_id
                    except json.JSONDecodeError:
                        pass
            
            # Case 3: Check for content array format: {"content": [{"text": "..."}]}
            if 'content' in parsed and isinstance(parsed['content'], list):
                content_list = parsed['content']
                if len(content_list) > 0:
                    first_content = content_list[0]
                    if isinstance(first_content, dict) and 'text' in first_content:
                        return first_content['text'], parsed.get('session_id')
    except json.JSONDecodeError:
        pass
    except Exception as e:
        print(f"[extract_text_from_analyst_response] Error: {e}")
    
    return None, None

# ─────────────────────────────────────────────────────────────────────────────
# Lambda Warm-Up endpoint
# Called silently when the BRD Analyst page opens to pre-warm Lambda containers
# before the user sends their first message.
# ─────────────────────────────────────────────────────────────────────────────










@app.get("/api/download-brd/{brd_id}")
async def download_brd(
    brd_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Download a generated BRD document as DOCX - fetches from S3 and converts to DOCX"""
    try:
        print(f"\n[DOWNLOAD-BRD] ========== START ==========")
        print(f"[DOWNLOAD-BRD] BRD ID: {brd_id}")
        print(f"[DOWNLOAD-BRD] User: {current_user.get('user_id', 'unknown')}")
        
        # Get S3 client
        s3_client = get_s3_client()
        bucket_name = S3_BUCKET_NAME
        
        # BRD is stored as: brds/{brd_id}/BRD_{brd_id}.txt
        s3_key_txt = f"brds/{brd_id}/BRD_{brd_id}.txt"
        
        print(f"[DOWNLOAD-BRD] Fetching TXT from s3://{bucket_name}/{s3_key_txt}")
        
        # Get the BRD TXT file from S3
        try:
            head_response = s3_client.head_object(Bucket=bucket_name, Key=s3_key_txt)
            print(f"[DOWNLOAD] ✅ File exists in S3! Size: {head_response.get('ContentLength', 0)} bytes")
        except ClientError as head_err:
            head_error_code = head_err.response.get('Error', {}).get('Code', 'Unknown')
            print(f"[DOWNLOAD] ⚠️  head_object failed: {head_error_code} - {head_err}")
        
        try:
            # Try to get BRD JSON structure first (preferred for DOCX conversion)
            # Check both possible naming conventions
            json_key = f"brds/{brd_id}/brd_structure.json"
            try:
                try:
                    json_response = s3_client.get_object(Bucket=bucket_name, Key=json_key)
                except ClientError as e:
                    if e.response.get('Error', {}).get('Code') == 'NoSuchKey':
                        json_key = f"brds/{brd_id}/BRD_{brd_id}.json"
                        json_response = s3_client.get_object(Bucket=bucket_name, Key=json_key)
                    else:
                        raise e
                
                # Read with explicit UTF-8 encoding and error handling
                json_body = json_response['Body'].read()
                print(f"[DOWNLOAD] Read {len(json_body)} bytes from JSON file")
                
                # Parse JSON
                try:
                    brd_data = json.loads(json_body)
                except json.JSONDecodeError as je:
                    print(f"[DOWNLOAD] ⚠️ Error parsing BRD JSON: {je}")
                    raise je

                # Convert JSON to DOCX
                docx_bytes = render_brd_json_to_docx(brd_data)
                
                print(f"[DOWNLOAD-BRD] ✅ Converted JSON to DOCX: {len(docx_bytes)} bytes")
                
                from fastapi.responses import Response
                return Response(
                    content=docx_bytes,
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={
                        "Content-Disposition": f"attachment; filename=BRD_{brd_id}.docx"
                    }
                )
            except ClientError as json_err:
                json_error_code = json_err.response.get('Error', {}).get('Code', '')
                if json_error_code == 'NoSuchKey':
                    print(f"[DOWNLOAD] ⚠️  BRD JSON not found, trying text file...")
                    # Fallback to text file and convert to DOCX
                    response = s3_client.get_object(Bucket=bucket_name, Key=s3_key_txt)
                    # Read with explicit UTF-8 encoding and error handling
                    text_body = response['Body'].read()
                    print(f"[DOWNLOAD] Read {len(text_body)} bytes from text file")
                    
                    try:
                        brd_text = text_body.decode('utf-8')
                    except UnicodeDecodeError as decode_err:
                        # Try with error handling
                        print(f"[DOWNLOAD] ⚠️ UTF-8 decode error: {decode_err}, trying with error replacement")
                        brd_text = text_body.decode('utf-8', errors='replace')
                    
                    print(f"[DOWNLOAD] ✅ Fetched BRD text from S3 ({len(brd_text)} chars)")
                    print(f"[DOWNLOAD] First 300 chars: {brd_text[:300]}")
                    print(f"[DOWNLOAD] Last 100 chars: {brd_text[-100:]}")
                    
                    # Check if text looks like it's in English (basic check)
                    english_chars = sum(1 for c in brd_text[:500] if c.isascii() and (c.isalpha() or c.isspace() or c in '.,;:!?()-'))
                    total_chars = min(500, len(brd_text))
                    if total_chars > 0:
                        english_ratio = english_chars / total_chars
                        print(f"[DOWNLOAD] English character ratio (first 500 chars): {english_ratio:.2%}")
                        if english_ratio < 0.5:
                            print(f"[DOWNLOAD] ⚠️ WARNING: Text may not be in English! Ratio: {english_ratio:.2%}")
                    
                    # Convert text to DOCX with proper markdown parsing
                    docx_content = render_brd_json_to_docx(brd_text)  # This handles string input
                    
                    print(f"[DOWNLOAD] ✅ Converted text to DOCX ({len(docx_content)} bytes)")
                    
                    from fastapi.responses import Response
                    return Response(
                        content=docx_content,
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        headers={
                            "Content-Disposition": f'attachment; filename="BRD_{brd_id}.docx"'
                        }
                    )
                else:
                    raise json_err
        except ClientError as e:
            error_code = e.response.get('Error', {}).get('Code', 'Unknown')
            error_message = e.response.get('Error', {}).get('Message', str(e))
            
            print(f"[DOWNLOAD] ❌ S3 ClientError: Code={error_code}, Message={error_message}")
            print(f"[DOWNLOAD] Full error response: {e.response}")
            
            # Check credentials
            try:
                creds_valid, creds_info = check_aws_credentials()
                if not creds_valid:
                    print(f"[DOWNLOAD] ⚠️  AWS credentials check failed: {creds_info}")
                    return JSONResponse(
                        status_code=403,
                        content={"error": f"AWS credentials invalid or expired. Please refresh credentials. Details: {creds_info}"}
                    )
                else:
                    print(f"[DOWNLOAD] ✅ AWS credentials valid: {creds_info.get('Account', 'Unknown')}")
            except Exception as cred_err:
                print(f"[DOWNLOAD] ⚠️  Could not check credentials: {cred_err}")
            
            # Check if it's an access denied error
            if error_code == 'AccessDeniedException' or '403' in str(error_code) or 'Access Denied' in error_message:
                return JSONResponse(
                    status_code=403,
                    content={"error": f"Access denied to S3 bucket '{bucket_name}'. Please check IAM permissions for s3:GetObject. Error: {error_message}"}
                )
            
            if error_code == 'NoSuchKey':
                print(f"[DOWNLOAD] ❌ BRD text file not found in S3: {s3_key_txt}")
                # Try to get BRD JSON structure and render it to text
                try:
                    print(f"[DOWNLOAD] Attempting to fetch BRD JSON structure from S3...")
                    json_key = f"brds/{brd_id}/brd_structure.json"
                    try:
                        try:
                            json_response = s3_client.get_object(Bucket=bucket_name, Key=json_key)
                        except ClientError as e:
                            if e.response.get('Error', {}).get('Code') == 'NoSuchKey':
                                json_key = f"brds/{brd_id}/BRD_{brd_id}.json"
                                json_response = s3_client.get_object(Bucket=bucket_name, Key=json_key)
                            else:
                                raise e
                        
                        # Read with explicit UTF-8 encoding and error handling
                        json_body = json_response['Body'].read()
                        try:
                            json_text = json_body.decode('utf-8')
                        except UnicodeDecodeError:
                            # Try with error handling
                            json_text = json_body.decode('utf-8', errors='replace')
                            print(f"[DOWNLOAD] ⚠️ Had to use error replacement for JSON decoding (fallback)")
                        
                        brd_json = json.loads(json_text)
                        print(f"[DOWNLOAD] ✅ Found BRD JSON structure, rendering to text...")
                        
                        # Render BRD JSON to text (see render_brd_json_to_text above).
                        brd_text = render_brd_json_to_text(brd_json)
                        
                        # Also save the text file for future downloads
                        try:
                            s3_put_object(
                                key=s3_key_txt,
                                body=brd_text,
                                content_type="text/plain",
                                bucket=bucket_name,
                            )
                            print(f"[DOWNLOAD] ✅ Saved rendered text file to S3 for future downloads")
                        except Exception as save_err:
                            print(f"[DOWNLOAD] ⚠️  Could not save text file: {save_err}")
                        
                        print(f"[DOWNLOAD] ✅ Rendered BRD from JSON ({len(brd_text)} chars)")
                        
                        # Convert to DOCX
                        docx_bytes = render_brd_json_to_docx(brd_json)
                        print(f"[DOWNLOAD] ✅ Generated DOCX ({len(docx_bytes)} bytes)")
                        
                        from fastapi.responses import Response
                        return Response(
                            content=docx_bytes,
                            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            headers={
                                "Content-Disposition": f'attachment; filename="BRD_{brd_id}.docx"'
                            }
                        )
                    except ClientError as json_err:
                        json_error_code = json_err.response.get('Error', {}).get('Code', '')
                        if json_error_code == 'NoSuchKey':
                            print(f"[DOWNLOAD] ❌ BRD JSON structure also not found")
                        else:
                            print(f"[DOWNLOAD] ❌ Error fetching JSON: {json_err}")
                except Exception as render_err:
                    print(f"[DOWNLOAD] ❌ Failed to render BRD from JSON: {render_err}")
                    import traceback
                    traceback.print_exc()
                
                return JSONResponse(
                    status_code=404,
                    content={"error": f"BRD {brd_id} not found in S3. Neither text file nor JSON structure found. The BRD may not have been saved yet."}
                )
            elif error_code == 'AccessDeniedException':
                return JSONResponse(
                    status_code=403,
                    content={"error": f"Access denied to S3 bucket. Please check AWS credentials and IAM permissions. Error: {error_message}"}
                )
            else:
                return JSONResponse(
                    status_code=500,
                    content={"error": f"S3 error ({error_code}): {error_message}"}
                )
        except Exception as e:
            print(f"[DOWNLOAD] ❌ Error fetching from S3: {e}")
            import traceback
            traceback.print_exc()
            return JSONResponse(status_code=500, content={
                "error": "Failed to retrieve BRD",
                "message": f"Error accessing S3: {str(e)}"
            })
            
    except Exception as e:
        error_msg = str(e)
        print(f"[DOWNLOAD-BRD] ERROR: {error_msg}")
        import traceback
        traceback.print_exc()
        
        return JSONResponse(status_code=500, content={
            "error": error_msg,
            "message": f"Error downloading BRD: {error_msg}"
        })


# -------------------------
# BRD Chat History Endpoint (my_agent memory)
# -------------------------

def _extract_clean_user_message(text: str) -> str:
    """Extract the clean user message from the enhanced context sent to the LLM.

    The frontend wraps user messages with section context in formats like:
      SECTION N: Title\\n\\n{content}\\n\\nUSER REQUEST: {actual message}\\n\\nIMPORTANT: ...
      BRD CONTEXT:\\n{content}\\n\\nUSER REQUEST: {actual message}
    This extracts just the {actual message} part for clean chat display.
    """
    import re

    # 1. Extract text after "USER REQUEST: " marker (strips the section context prefix)
    marker = "USER REQUEST: "
    idx = text.find(marker)
    if idx != -1:
        clean = text[idx + len(marker):]
    else:
        clean = text

    # 2. Strip trailing IMPORTANT instruction block — use regex to handle any whitespace (\n, \r\n, etc.)
    clean = re.split(r'\s+IMPORTANT:\s+The user is currently viewing', clean, maxsplit=1)[0]

    return clean.strip()


@app.get("/api/brd-history/{session_id}")
async def get_brd_chat_history(
    session_id: str,
    project_id: str = None,
    current_user: dict = Depends(get_current_user)
):
    """Get conversation history for BRD chat from AgentCore Memory.

    Uses the same memory the Lambda writes to (AGENTCORE_MEMORY_ID),
    mirroring the working /api/analyst-history pattern.
    """
    try:
        print(f"\n[BRD-HISTORY] Retrieving history from AgentCore Memory for session: {session_id}")

        agentcore_client = get_agent_core_client()
        memory_id = AGENTCORE_MEMORY_ID
        actor_id = AGENTCORE_ACTOR_ID

        print(f"[BRD-HISTORY] Query params: memoryId={memory_id}, sessionId={session_id}, actorId={actor_id}")

        messages = []

        try:
            response = agentcore_client.list_events(
                memoryId=memory_id,
                sessionId=session_id,
                actorId=actor_id,
                includePayloads=True,
                maxResults=99
            )

            events = response.get("events", [])
            print(f"[BRD-HISTORY] AgentCore returned {len(events)} events")

            # Sort events by (eventTimestamp, eventId) oldest first.
            # list_events may return in undefined order. Using eventId as
            # a secondary key ensures correct ordering when timestamps
            # are identical or have low precision.
            def _event_sort_key(e):
                ts = e.get("eventTimestamp")
                if ts is None:
                    ts_str = ""
                else:
                    ts_str = ts.isoformat() if hasattr(ts, "isoformat") else str(ts)
                eid = e.get("eventId", "")
                return (ts_str, eid)
            events = sorted(events, key=_event_sort_key)

            for event in events:
                payload_list = event.get("payload", [])
                for payload_item in payload_list:
                    conv_data = payload_item.get("conversational")
                    if not conv_data:
                        continue
                    text_content = conv_data.get("content", {}).get("text")
                    if not text_content:
                        continue
                    role = conv_data.get("role", "assistant").lower()
                    if role == "user":
                        # Strip the enhanced section context; keep only the actual user message
                        clean_text = _extract_clean_user_message(text_content)
                        messages.append({
                            "role": "user",
                            "content": clean_text,
                            "isBot": False
                        })
                    elif role == "assistant":
                        messages.append({
                            "role": "assistant",
                            "content": text_content,
                            "isBot": True
                        })

            # Already sorted oldest-first above — no need to reverse

        except Exception as e:
            print(f"[BRD-HISTORY] AgentCore Memory query failed: {e}")
            import traceback
            traceback.print_exc()

        print(f"[BRD-HISTORY] Returning {len(messages)} messages")

        return JSONResponse(content={
            "messages": messages,
            "session_id": session_id,
            "count": len(messages)
        })

    except Exception as e:
        print(f"[BRD-HISTORY] ERROR: {e}")
        import traceback
        traceback.print_exc()
        return JSONResponse(status_code=500, content={
            "error": str(e),
            "messages": []
        })

# -------------------------
# Access Control Endpoints
# -------------------------

@app.get("/api/brd/access")
async def check_brd_access(current_user: dict = Depends(get_current_user)):
    """Check if current user has BRD access"""
    user_id = current_user["user_id"]
    has_access = check_brd_access_via_agentcore(user_id)
    
    return JSONResponse(content={
        "has_access": has_access,
        "user_id": user_id,
        "email": current_user["email"]
    })

@app.get("/api/user/info")
async def get_user_info(current_user: dict = Depends(get_current_user)):
    """Get current user information including group-based module access.

    Source of truth for RBAC — frontend calls this after login instead of
    computing modules from idTokenClaims.groups, so that:
      1. Groups-claim overage (>200 groups) is handled via Microsoft Graph
         inside extract_user_groups().
      2. Every authenticated user is recorded in the DB (even those who end
         up with no allowed modules and see AccessDenied).
    """
    user_id = current_user["user_id"]
    email = current_user["email"]
    name = current_user.get("name", "")

    try:
        create_or_update_user(user_id, email, name)
    except Exception as e:
        print(f"[USER_INFO] Failed to upsert user row for {email}: {e}")

    identity_arn = get_user_identity_arn(user_id)
    return JSONResponse(content={
        "user_id": user_id,
        "email": email,
        "name": name,
        "identity_arn": identity_arn,
        "groups": current_user.get("groups", []),
        "allowed_modules": current_user.get("allowed_modules", []),
    })


@app.get("/api/user/me/usage")
async def get_my_usage(current_user: dict = Depends(get_current_user)):
    """Return the current user's own usage row (last_login + cumulative tokens)
    plus per-module rollup and recent events."""
    from db_helper import get_user_usage, get_user_module_rollup, get_user_recent_events
    user_id = current_user["user_id"]
    row = get_user_usage(user_id)
    modules = get_user_module_rollup(user_id)
    recent_events = get_user_recent_events(user_id, limit=20)
    if not row:
        return JSONResponse(content={
            "user_id": user_id,
            "email": current_user.get("email"),
            "name": current_user.get("name"),
            "created_at": None,
            "last_login": None,
            "token_usage": 0,
            "access_role": "NONE",
            "modules": modules,
            "recent_events": recent_events,
        })
    return JSONResponse(content={
        "user_id": row["id"],
        "email": row["email"],
        "name": row["name"],
        "created_at": row["created_at"].isoformat() if row.get("created_at") else None,
        "last_login": row["last_login"].isoformat() if row.get("last_login") else None,
        "token_usage": int(row.get("token_usage") or 0),
        "access_role": row.get("access_role") or "NONE",
        "modules": modules,
        "recent_events": recent_events,
    })


@app.get("/api/users/usage")
async def get_organization_usage(current_user: dict = Depends(get_current_user)):
    """Roster of all users with token usage, last login, per-module rollup,
    and recent events. Visible to any authenticated user.

    Per-user `recent_events` is capped at 10 to keep payload size bounded;
    `/api/user/me/usage` returns 20 for the caller's own drill-down.
    """
    from db_helper import (
        list_all_users_usage,
        get_user_module_rollup,
        get_user_recent_events,
    )
    rows = list_all_users_usage()
    users_payload = []
    for r in rows:
        uid = r["id"]
        users_payload.append({
            "user_id": uid,
            "email": r["email"],
            "name": r["name"],
            "created_at": r["created_at"].isoformat() if r.get("created_at") else None,
            "last_login": r["last_login"].isoformat() if r.get("last_login") else None,
            "token_usage": int(r.get("token_usage") or 0),
            "is_active": bool(r.get("is_active", True)),
            "access_role": r.get("access_role") or "NONE",
            "modules": get_user_module_rollup(uid),
            "recent_events": get_user_recent_events(uid, limit=10),
        })
    return JSONResponse(content={
        "users": users_payload,
        "total_users": len(rows),
        "total_tokens": sum(int(r.get("token_usage") or 0) for r in rows),
    })


@app.get("/api/support/user-guide")
async def get_support_user_guide(current_user: dict = Depends(get_current_user)):
    """Return the user guide HTML extracted from the MHTML .doc in S3."""
    import email as email_lib
    import base64 as b64

    s3_client = get_s3_client()
    try:
        response = s3_client.get_object(
            Bucket=S3_BUCKET_NAME,
            Key="support/SDLC_Orchestrator_Userguide_ForTesting.doc",
        )
        raw = response["Body"].read().decode("utf-8", errors="ignore")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to load user guide from S3: {e}")

    msg = email_lib.message_from_string(raw)
    image_map: dict[str, str] = {}
    html_content = ""

    for part in msg.walk():
        ct = part.get_content_type()
        payload = part.get_payload(decode=True)
        if not payload:
            continue
        if "html" in ct:
            html_content = payload.decode("utf-8", errors="ignore")
            continue
        if ct.startswith("multipart/"):
            continue
        mime = ct
        if ct == "application/octet-stream":
            if payload[:8] == b"\x89PNG\r\n\x1a\n":
                mime = "image/png"
            elif payload[:2] == b"\xff\xd8":
                mime = "image/jpeg"
            elif payload[:4] == b"GIF8":
                mime = "image/gif"
            else:
                mime = "image/png"
        data_uri = f"data:{mime};base64,{b64.b64encode(payload).decode()}"
        loc = part.get("Content-Location", "")
        if loc:
            image_map[loc] = data_uri
            fname = loc.rsplit("/", 1)[-1] if "/" in loc else loc
            if fname:
                image_map[fname] = data_uri
        cid = part.get("Content-ID", "").strip("<>")
        if cid:
            image_map[f"cid:{cid}"] = data_uri

    if not html_content:
        raise HTTPException(status_code=500, detail="Could not extract HTML from user guide")

    for ref in sorted(image_map, key=len, reverse=True):
        html_content = html_content.replace(ref, image_map[ref])

    # Add IDs to headings so the frontend can scroll to specific sections
    import re as _re

    def _slugify(text: str) -> str:
        clean = _re.sub(r"<[^>]+>", "", text).strip()
        return _re.sub(r"[^a-z0-9]+", "-", clean.lower()).strip("-")

    def _add_heading_ids(html: str) -> str:
        def _replacer(m):
            tag = m.group(1)
            attrs = m.group(2)
            content = m.group(3)
            slug = _slugify(content)
            if slug:
                return f"<{tag}{attrs} id=\"{slug}\">{content}</{tag}>"
            return m.group(0)
        return _re.sub(r"<(h[1-4])([^>]*)>(.*?)</\1>", _replacer, html, flags=_re.DOTALL | _re.IGNORECASE)

    html_content = _add_heading_ids(html_content)

    style_block = """
    <style>
      body, .WordSection1 {
        font-family: 'Segoe UI', -apple-system, BlinkMacSystemFont, Arial, sans-serif;
        line-height: 1.9; color: #1e293b; font-size: 19px;
      }
      h1 { color: #0f172a; font-size: 38px; font-weight: 700; margin: 1.8em 0 0.8em; padding-bottom: 0.4em; border-bottom: 2px solid #e2e8f0; }
      h2 { color: #1e40af; font-size: 30px; font-weight: 700; margin: 1.8em 0 0.6em; padding-bottom: 0.3em; border-bottom: 1px solid #e2e8f0; }
      h3 { color: #1e40af; font-size: 24px; font-weight: 600; margin: 1.5em 0 0.5em; }
      h4 { color: #334155; font-size: 20px; font-weight: 600; margin: 1.2em 0 0.4em; }
      p { margin: 0.8em 0; color: #374151; font-size: 17px; }
      img { max-width: 50%; height: auto; border-radius: 10px; margin: 24px 0; display: block; box-shadow: 0 4px 16px rgba(0,0,0,0.12); border: 1px solid #e5e7eb; }
      table { border-collapse: collapse; width: 100%; margin: 1.2em 0; border-radius: 8px; overflow: hidden; box-shadow: 0 1px 3px rgba(0,0,0,0.08); }
      th { background: #f1f5f9; font-weight: 600; color: #1e293b; text-align: left; }
      td, th { border: 1px solid #e2e8f0; padding: 10px 14px; font-size: 14px; }
      tr:nth-child(even) { background: #f8fafc; }
      a { color: #2563eb; text-decoration: none; font-weight: 500; }
      a:hover { text-decoration: underline; color: #1d4ed8; }
      ul, ol { margin: 0.6em 0; padding-left: 1.8em; }
      li { margin: 0.4em 0; color: #374151; }
      code { background: #f1f5f9; padding: 2px 6px; border-radius: 4px; font-size: 13px; color: #be185d; }
      pre { background: #1e293b; color: #e2e8f0; padding: 16px; border-radius: 8px; overflow-x: auto; font-size: 13px; }
      blockquote { border-left: 4px solid #3b82f6; margin: 1em 0; padding: 0.5em 1em; background: #eff6ff; border-radius: 0 6px 6px 0; }
    </style>
    """
    html_content = style_block + html_content
    return JSONResponse(content={"html": html_content})


@app.post("/api/admin/grant-brd-access")
async def grant_brd_access(
    target_user_id: str = Form(...),
    current_user: dict = Depends(get_current_user)
):
    """Grant BRD access to a user"""
    success = grant_brd_access_via_agentcore(target_user_id)
    return JSONResponse(content={"success": success, "user_id": target_user_id})

@app.post("/api/admin/revoke-brd-access")
async def revoke_brd_access(
    target_user_id: str = Form(...),
    current_user: dict = Depends(get_current_user)
):
    """Revoke BRD access from a user"""
    success = revoke_brd_access_via_agentcore(target_user_id)
    return JSONResponse(content={"success": success, "user_id": target_user_id})


# -------------------------
# BRD Read APIs (S3-backed)
# -------------------------

def _load_brd_structure_from_s3(brd_id: str) -> dict:
    """Load the latest BRD structure JSON from S3."""
    s3_client = get_s3_client()
    bucket_name = S3_BUCKET_NAME
    key = f"brds/{brd_id}/brd_structure.json"
    try:
        response = s3_client.get_object(Bucket=bucket_name, Key=key)
        body = response["Body"].read()
        return json.loads(body)
    except ClientError as e:
        code = e.response.get("Error", {}).get("Code", "Unknown")
        raise HTTPException(status_code=404, detail=f"BRD structure not found in S3 ({code}): {key}")
    except json.JSONDecodeError as e:
        raise HTTPException(status_code=500, detail=f"BRD structure JSON is invalid: {str(e)}")


def _is_doc_title_section(title: str) -> bool:
    """Detect if the first section is a document title (not a real BRD section).

    Real BRD sections always start with a number like "1. Document Overview".
    The LLM sometimes prepends a project-name section without a number prefix;
    this helper identifies those so ``_iter_user_sections`` can skip them.
    """
    t = (title or "").strip()
    if not t:
        return False
    # If the first section starts with a digit it is a real numbered section.
    if re.match(r"^\d+\.", t):
        return False
    # Everything else (project names, headings without numbers) is a doc title.
    return True


def _iter_user_sections(brd_data: dict):
    """Yield (user_section_number, array_index, title, section_dict) for user-visible sections."""
    sections = brd_data.get("sections", []) if isinstance(brd_data, dict) else []
    if not sections:
        return

    start_idx = 0
    if sections and _is_doc_title_section(sections[0].get("title", "")):
        start_idx = 1

    user_num = 1
    for idx in range(start_idx, len(sections)):
        sec = sections[idx]
        title = (sec.get("title", "") or "").strip()
        title_lower = title.lower()

        # Skip ALL subsections whose title starts with "#"
        # These are sub-headers within a parent section (e.g. "# User Story 1: ...",
        # "# In Scope", "# Out of Scope", "# Acronyms and Abbreviations", "# Appendix")
        # and should not appear as top-level BRD sections.
        if title.startswith("#"):
            continue

        yield user_num, idx, title, sec
        user_num += 1


def _get_user_section_by_number(brd_data: dict, section_number: int) -> dict:
    for user_num, idx, title, sec in _iter_user_sections(brd_data) or []:
        if user_num == section_number:
            return {"array_index": idx, "title": title, "section": sec}
    raise HTTPException(status_code=404, detail=f"Section {section_number} not found")


def _render_section_to_markdown(section_number: int, title: str, section: dict, brd_data: dict) -> str:
    sections = brd_data.get("sections", [])
    content_blocks = list(section.get("content", []) or [])

    # Scope (section 5): merge content from "# In Scope" and "# Out of Scope" if main section is empty
    if (not content_blocks or len(content_blocks) == 0) and "scope" in (title or "").lower():
        merged = []
        # Find the array index of the scope section and look ahead for subsections
        scope_idx = None
        for user_num, idx, t, _sec in _iter_user_sections(brd_data) or []:
            if user_num == section_number:
                scope_idx = idx
                break
        if scope_idx is not None:
            for i in (1, 2):
                sub_idx = scope_idx + i
                if sub_idx < len(sections):
                    sub = sections[sub_idx]
                    sub_title = (sub.get("title", "") or "").lower()
                    if "# in scope" in sub_title:
                        merged.append({"type": "paragraph", "text": "### In Scope"})
                        merged.extend(sub.get("content", []) or [])
                    elif "# out of scope" in sub_title:
                        merged.append({"type": "paragraph", "text": "### Out of Scope"})
                        merged.extend(sub.get("content", []) or [])
        if merged:
            content_blocks = merged

    title_clean = re.sub(r"^\d+\.\s*", "", title or "").strip() or "Untitled"
    md = f"## {section_number}. {title_clean}\n\n"

    for block in content_blocks:
        if not isinstance(block, dict):
            continue
        block_type = block.get("type")
        if block_type == "paragraph":
            md += (block.get("text", "") or "") + "\n\n"
        elif block_type == "bullet":
            for item in block.get("items", []) or []:
                md += f"- {item}\n"
            md += "\n"
        elif block_type == "table":
            rows = block.get("rows", []) or []
            for row in rows:
                md += "| " + " | ".join(str(cell) for cell in row) + " |\n"
            md += "\n"

    return md.strip() + "\n"


@app.get("/api/brd/{brd_id}/structure")
async def api_get_brd_structure(
    brd_id: str,
    current_user: dict = Depends(get_current_user),
):
    """Return the latest structured BRD JSON from S3 (source of truth for sections)."""
    brd_data = _load_brd_structure_from_s3(brd_id)
    return JSONResponse(content={"brd_id": brd_id, "brd": brd_data})


@app.get("/api/brd/{brd_id}/sections")
async def api_list_brd_sections(
    brd_id: str,
    current_user: dict = Depends(get_current_user),
):
    """Return user-visible section numbers + titles (use this to build tabs)."""
    brd_data = _load_brd_structure_from_s3(brd_id)
    sections = [{"number": n, "title": re.sub(r"^\\d+\\.\\s*", "", t).strip() or t} for n, _idx, t, _sec in (_iter_user_sections(brd_data) or [])]
    return JSONResponse(content={"brd_id": brd_id, "sections": sections})


@app.get("/api/brd/{brd_id}/section/{section_number}")
async def api_get_brd_section(
    brd_id: str,
    section_number: int,
    current_user: dict = Depends(get_current_user),
):
    """
    Return the latest version of a single section from S3.

    This is the safest way to power section-tabs: it always reads the newest `brd_structure.json`.
    """
    brd_data = _load_brd_structure_from_s3(brd_id)
    found = _get_user_section_by_number(brd_data, section_number)
    title = found["title"]
    section = found["section"]
    markdown = _render_section_to_markdown(section_number, title, section, brd_data)
    return JSONResponse(content={
        "brd_id": brd_id,
        "section_number": section_number,
        "title": title,
        "section": section,
        "markdown": markdown,
    })


# -------------------------
# Internal token-usage callback
# -------------------------
# Used by Lambdas and AgentCore agents (which can't reach RDS directly) to
# report LLM token consumption back to the backend. API-key authenticated.

class RecordTokensRequest(BaseModel):
    user_id: str
    tokens: int
    source: Optional[str] = None  # e.g. "lambda_brd_generator", "pm_agent"


@app.post("/api/internal/record-tokens", status_code=204)
async def record_tokens(
    body: RecordTokensRequest,
    x_api_key: str = Header(alias="X-API-Key"),
):
    """Increment users.token_usage from Lambda/agent call sites that can't talk
    to RDS directly. Validates X-API-Key against INTERNAL_API_KEYS env."""
    from routers.internal_utils import validate_api_key
    from db_helper import increment_user_token_usage as _bump

    validate_api_key(x_api_key)

    if body.tokens <= 0 or not body.user_id:
        return  # 204 — silently no-op for invalid inputs

    try:
        _bump(body.user_id, body.tokens)
        print(f"[record-tokens] user={body.user_id} tokens={body.tokens} source={body.source or '?'}")
    except Exception as e:
        # Log but don't surface — token accounting must never break callers
        print(f"[record-tokens] FAILED user={body.user_id} tokens={body.tokens}: {e}")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
